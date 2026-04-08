"""
Sistema de Gestión de Dotación APS
Servicio de Salud Metropolitano Central — 2026
CRUD local con Excel en tiempo real.
"""
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from pathlib import Path
from copy import deepcopy
import openpyxl
from openpyxl.utils import get_column_letter
import warnings, gc, os, sys, shutil, tempfile, sqlite3, io, hmac, requests as _req
from datetime import datetime, date, time as dt_time
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, DataReturnMode, JsCode

warnings.filterwarnings("ignore")

# Directorio de trabajo temporal para copias del Excel
_WORK_DIR = Path(tempfile.gettempdir()) / "app_dotacion_ssmc"
_WORK_DIR.mkdir(exist_ok=True)

# SQLite: Base de datos local para persistencia de la revisión
_DB_DIR = Path(__file__).resolve().parent / "data"
_DB_DIR.mkdir(exist_ok=True)
_DB_PATH = _DB_DIR / "revision_dotacion.db"

# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Dotación APS | SSMC",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Paleta corporativa (Evidant: verde oscuro + lima neón) ──
C_PRIMARY   = "#0A2E1F"
C_ACCENT    = "#4ADE80"
C_SUCCESS   = "#22C55E"
C_WARNING   = "#FACC15"
C_DANGER    = "#EF4444"
C_BG_CARD   = "#F0FDF4"
C_TEXT      = "#14352A"

# ─── CSS ─────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

.main-header {
    background: linear-gradient(135deg, #0A2E1F 0%, #145A38 60%, #1A7A4A 100%);
    color: white; padding: 1.2rem 1.5rem; border-radius: 12px;
    margin-bottom: 1rem; box-shadow: 0 4px 20px rgba(10,46,31,.35);
    border-bottom: 3px solid #4ADE80;
}
.main-header h1 { color: white; margin: 0; font-size: 1.5rem; font-weight: 700; }
.main-header p  { color: #86EFAC; margin: .15rem 0 0; font-size: .85rem; }

.kpi-row { display: flex; gap: 12px; flex-wrap: wrap; margin-bottom: 1rem; }
.kpi-card {
    flex: 1 1 160px; background: white; border-radius: 10px;
    padding: .85rem 1rem; text-align: center;
    border-top: 4px solid #4ADE80;
    box-shadow: 0 2px 8px rgba(10,46,31,.08);
    min-width: 140px;
}
.kpi-card.green  { border-top-color: #22C55E; }
.kpi-card.amber  { border-top-color: #FACC15; }
.kpi-card.red    { border-top-color: #EF4444; }
.kpi-value { font-size: 1.65rem; font-weight: 700; color: #0A2E1F; line-height: 1; }
.kpi-label { font-size: .72rem; color: #555; margin-top: .2rem; text-transform: uppercase;
             letter-spacing: .5px; }
.kpi-sub   { font-size: .7rem; color: #888; }

.section-title {
    font-size: 1rem; font-weight: 600; color: #0A2E1F;
    border-left: 4px solid #4ADE80; padding-left: .6rem;
    margin: 1.2rem 0 .6rem;
}

div[data-testid="stDataEditor"] { border-radius: 8px; overflow: hidden; }
div[data-testid="stMetric"] {
    background: white; border-radius: 8px; padding: .75rem;
    box-shadow: 0 1px 4px rgba(10,46,31,.06);
}

.save-banner {
    background: #DCFCE7; border-left: 5px solid #22C55E;
    padding: .6rem 1rem; border-radius: 6px; font-size: .85rem;
    color: #166534; margin: .5rem 0;
}
.warn-banner {
    background: #FEF9C3; border-left: 5px solid #FACC15;
    padding: .6rem 1rem; border-radius: 6px; font-size: .85rem;
    color: #854D0E; margin: .5rem 0;
}

/* Tabs styling */
button[data-baseweb="tab"] {
    font-weight: 600 !important; font-size: .85rem !important;
}
/* Accent overrides for Streamlit widgets */
.stButton > button {
    background-color: #0A2E1F; color: #4ADE80; border: 1px solid #4ADE80;
    font-weight: 600; border-radius: 8px; transition: all .2s;
}
.stButton > button:hover {
    background-color: #145A38; color: white; border-color: #4ADE80;
    box-shadow: 0 0 12px rgba(74,222,128,.3);
}
.stButton > button:active { background-color: #0A2E1F; }

/* Sidebar accent */
section[data-testid="stSidebar"] {
    border-right: 3px solid #4ADE80;
}
/* Compact multiselect labels for inline filters */
.stMultiSelect label, .stSelectbox label, .stTextInput label {
    font-size: .78rem !important; font-weight: 600 !important;
    color: #0A2E1F !important; margin-bottom: 0 !important;
}
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stMultiSelect label {
    color: #0A2E1F; font-weight: 600;
}

/* ── Encabezado de tabla: fondo verde oscuro con texto claro ── */
div[data-testid="stDataEditor"] th,
div[data-testid="stDataFrame"] th {
    background: linear-gradient(135deg, #0A2E1F, #145A38) !important;
    color: #4ADE80 !important;
    font-weight: 700 !important;
    font-size: .78rem !important;
    letter-spacing: .3px;
    border-bottom: 2px solid #4ADE80 !important;
    position: sticky !important;
    top: 0 !important;
    z-index: 2 !important;
}
/* Glide header cells (data_editor uses glide-data-grid) */
div[data-testid="stDataEditor"] [class*="header"],
div[data-testid="stDataEditor"] [role="columnheader"] {
    background: #0A2E1F !important;
    color: #4ADE80 !important;
}

/* ── Data editor: bordes limpios ── */
div[data-testid="stDataEditor"] {
    border: 2px solid #0A2E1F;
    border-radius: 8px;
    overflow: hidden;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────
# UTILIDADES EXCEL
# ─────────────────────────────────────────────────────────────
SHEET_MAIN   = "TIT-CONTRATA MARZO 2026 (2)"
SHEET_HORAS  = "Horas Indirectas"
SHEET_DOT    = "DOT IDEAL"

# Mapeo cargo → columna DOT IDEAL (exhaustivo, por similitud)
CARGO_DOT_MAP = {
    # Médicos (todas las variantes)
    "MEDICO CIRUJANO":                     "MEDICOS",
    "MEDICO CIRUJANO EDF":                 "MEDICOS",
    "MEDICO GENERAL":                      "MEDICOS",
    "MEDICO EDF ART 8ª":                   "MEDICOS",
    "MEDICO EDF ART 9ª":                   "MEDICOS",
    "MEDICO BECADO PRIMARIO":              "MEDICOS",
    "MEDICO EN COMISION DE ESTUDIOS":      "MEDICOS",
    "MEDICO PERIODO ASISTENCIAL OBLIGATORIO": "MEDICOS",
    "MEDICINA FAMILIAR":                   "MEDICOS",
    "PEDIATRIA":                           "MEDICOS",
    "M.C.OBSTETRICIA Y GINECOLOGIA":       "MEDICOS",
    "CIRUJANO DENTISTA":                   "ODONT.",
    # Odontólogos
    "ODONTOLOGO":                          "ODONT.",
    "ODONTOLOGO (A)":                      "ODONT.",
    "ODONTOLOGO(A)":                       "ODONT.",
    # Enfermeras
    "ENFERMERA (O)":                       "EU",
    "ENFERMERA(O)":                        "EU",
    # Nutricionista
    "NUTRICIONISTA":                       "NUTRI",
    # Matronas
    "MATRONA":                             "MATRONA",
    "MATRON (A)":                          "MATRONA",
    "MATRON(A)":                           "MATRONA",
    # Psicólogos
    "PSICOLOGO":                           "PS",
    "PSICOLOGO (A)":                       "PS",
    "PSICOLOGO(A)":                        "PS",
    # Trabajo social
    "ASISTENTE SOCIAL":                    "TS",
    "TRABAJADOR (A) SOCIAL":               "TS",
    # Químico farmacéutico
    "QUIMICO FARMACEUTICO":                "QF",
    "QUIMICO FARMACEUTICO (A)":            "QF",
    # TENS (todas las variantes)
    "TECNICOS DE NIVEL SUPERIOR EN ENFERMERIA":   "TENS",
    "TECNICOS DE NIVEL SUPERIOR EN ODONTOLOGIA":  "TENS",
    "TECNICO EN ENFERMERIA":               "TENS",
    "TECNICO EN ODONTOLOGIA":              "TENS",
    "TECNICO EN OBSTETRICIA":              "TENS",
    "TECNICO PARAMEDICO":                  "TENS",
    "TECNICO EN LABORATORIO":              "TENS",
    "TECNICO EN IMAGENOLOGIA":             "TENS",
    "AUXILIAR DE ENFERMERIA":              "TENS",
    "AUXILIAR DE ODONTOLOGIA":             "TENS",
    "AUXILIAR DE RADIOLOGIA":              "TENS",
    # Administrativos
    "ADMINISTRATIVO":                      "ADM",
    "COORDINADOR ADMINISTRATIVO":          "ADM",
    "CONTADOR":                            "ADM",
}


def _normalize_cesfam(name: str) -> str:
    """Normaliza nombre CESFAM para matching entre tablas.
    Extrae la palabra clave principal del nombre."""
    import unicodedata
    s = name.strip().upper()
    # Quitar acentos
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    # Quitar prefijos comunes
    for prefix in ("CESFAM ", "CES "):
        if s.startswith(prefix):
            s = s[len(prefix):]
    # Normalizar variantes conocidas
    s = s.replace("°", "").replace("º", "").replace("Nº", "N").replace("N ", "N")
    s = s.replace(".", " ").replace("  ", " ").strip()
    return s


# Tabla de homologación CESFAM: nombre normalizado → nombre canónico en DOT IDEAL
_CESFAM_HOMOLOG = {
    "MAIPU":              "CESFAM MAIPÚ",
    "N1":                 "CESFAM N°1",
    "N 1":                "CESFAM N°1",
    "N5":                 "CESFAM N°5",
    "N 5":                "CESFAM N°5",
    "A M JURICIC":        "CESFAM JURICIC",
    "JURICIC":            "CESFAM JURICIC",
    "AHUES":              "CESFAM AHUÉS",
    "AHUÉS":              "CESFAM AHUÉS",
    "N VOULLIEME":        "CESFAM VOULLIEME",
    "VOULLIEME":          "CESFAM VOULLIEME",
    "SOFIA PINCHEIRA":    "CESFAM PINCHEIRA",
    "PINCHEIRA":          "CESFAM PINCHEIRA",
    "LAS MERCEDES":       "CESFAM LAS MERCEDES",
    "CHUCHUNCO":          "CESFAM CHUCHUNCO",
    "PVI":                "CESFAM PVI",
    "PADRE VICENTE IRARRAZAVAL": "CESFAM PVI",
}


def match_cesfam_to_dot(cesfam_main: str, df_dot) -> str | None:
    """Busca el CESFAM de la hoja principal en df_dot usando homologación."""
    norm = _normalize_cesfam(cesfam_main)
    # 1) Coincidencia directa por homologación
    canon = _CESFAM_HOMOLOG.get(norm)
    if canon:
        # Buscar en df_dot
        for dot_cesfam in df_dot["CESFAM"].dropna().unique():
            if str(dot_cesfam).strip().upper() == canon.upper():
                return str(dot_cesfam).strip()
    # 2) Coincidencia por substring
    for dot_cesfam in df_dot["CESFAM"].dropna().unique():
        dot_norm = _normalize_cesfam(str(dot_cesfam))
        if norm == dot_norm or norm in dot_norm or dot_norm in norm:
            return str(dot_cesfam).strip()
    # 3) Coincidencia por palabra clave más larga
    norm_words = set(norm.split())
    best, best_score = None, 0
    for dot_cesfam in df_dot["CESFAM"].dropna().unique():
        dot_words = set(_normalize_cesfam(str(dot_cesfam)).split())
        common = norm_words & dot_words
        if len(common) > best_score:
            best_score = len(common)
            best = str(dot_cesfam).strip()
    if best_score >= 1 and best:
        return best
    return None

# Columnas principales de la hoja (no son encomendaciones individuales)
_KNOWN_MAIN_COLS = {
    "REVISADO", "TIPO", "RUT", "DV", "CESFAM", "NOMBRE PROFESIONAL",
    "DESCRIPCIÓN DE PLANTA", "CARGO", "UNIDAD DE DESEMPEÑO",
    "ENCOMENDACIONES", "OBSERVACIONES", "Horas por contrato",
    "Horas Totales", "Total Descuentos semanal (horas)", "Total Horas Clínicas",
}


def _list_excel_sheets(path: str) -> list[str]:
    """Lista los nombres de hojas de un archivo Excel sin cargarlo completo."""
    try:
        readable = _ensure_readable(Path(path))
        wb = openpyxl.load_workbook(readable, read_only=True, data_only=True)
        names = wb.sheetnames
        wb.close()
        return names
    except Exception:
        return []


def find_excel_file() -> Path | None:
    """Busca el Excel en ubicaciones conocidas."""
    candidates = [
        Path(r"C:\Users\DAP\OneDrive - SUBSECRETARIA DE SALUD PUBLICA\Escritorio 2024\GESTION\REVISIÓN DOTACIONES\REVISION DOTACION\Dotación Marzo 2026-02-04-2026.xlsx"),
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def _ensure_readable(path: Path) -> str:
    """Devuelve ruta legible. Si el archivo está bloqueado, crea copia temp."""
    try:
        with open(str(path), "rb") as f:
            f.read(4)
        return str(path)
    except PermissionError:
        tmp_path = _WORK_DIR / "dotacion_work.xlsx"
        try:
            import subprocess
            result = subprocess.run(
                ["powershell", "-Command",
                 f'Copy-Item -Path "{path}" -Destination "{tmp_path}" -Force'],
                capture_output=True, timeout=30
            )
        except Exception:
            shutil.copy2(str(path), str(tmp_path))
        if tmp_path.exists():
            return str(tmp_path)
        raise PermissionError(f"No se puede acceder al archivo: {path}")


@st.cache_data(ttl=300)
def load_excel(path: str, _mtime: float, sheet_main: str | None = None):
    """Carga las 3 hojas como DataFrames. Cache invalidado por mtime."""
    _sheet = sheet_main or SHEET_MAIN
    # ── Hoja principal ──
    df_main = pd.read_excel(path, sheet_name=_sheet, header=1, dtype=str)
    # Normalizar nombres de columna (pueden tener saltos de línea por celdas merge)
    clean_cols = []
    for c in df_main.columns:
        c_str = str(c).strip().split("\n")[0].strip()
        clean_cols.append(c_str)
    df_main.columns = clean_cols
    # Identificar columna NOMBRE (puede tener sufijo de fecha)
    for i, c in enumerate(df_main.columns):
        if c.upper().startswith("NOMBRE PROFESIONAL"):
            df_main.columns = list(df_main.columns[:i]) + ["NOMBRE PROFESIONAL"] + list(df_main.columns[i+1:])
            break
    # Limpiar filas vacías
    id_col = "RUT" if "RUT" in df_main.columns else df_main.columns[0]
    name_col = "NOMBRE PROFESIONAL" if "NOMBRE PROFESIONAL" in df_main.columns else df_main.columns[1]
    df_main = df_main.dropna(subset=[id_col, name_col], how="all")
    # Desduplicar columnas repetidas (ej: "Encargado tecnovigilancia" x2)
    seen = {}
    new_cols = []
    for c in df_main.columns:
        if c in seen:
            seen[c] += 1
            new_cols.append(f"{c}_{seen[c]}")
        else:
            seen[c] = 0
            new_cols.append(c)
    df_main.columns = new_cols
    # Eliminar columnas basura (nombre numérico o vacío)
    junk = [c for c in df_main.columns if c in ("0", "None", "") or (c.replace(".", "").isdigit())]
    if junk:
        df_main = df_main.drop(columns=junk)
    # Columna REVISADO si no existe
    if "REVISADO" not in df_main.columns:
        df_main.insert(0, "REVISADO", "")
    # Asegurar que REVISADO nunca tenga None
    df_main["REVISADO"] = df_main["REVISADO"].fillna("")

    # ── Consolidar columnas de encomendaciones individuales (SI/vacío) ──
    # Las columnas que coinciden con Items de Horas Indirectas se fusionan
    # en una sola columna ENCOMENDACIONES (texto separado por " | ")
    _main_keep = {"REVISADO", "TIPO", "RUT", "DV", "CESFAM", "NOMBRE PROFESIONAL",
                  "DESCRIPCIÓN DE PLANTA", "CARGO", "UNIDAD DE DESEMPEÑO",
                  "ENCOMENDACIONES", "OBSERVACIONES", "Horas por contrato",
                  "Horas Totales", "Total Descuentos semanal (horas)", "Total Horas Clínicas"}
    encom_individual_cols = [c for c in df_main.columns if c not in _main_keep]

    # Construir la lista combinada para ENCOMENDACIONES a partir de marcas SI
    combined_encom = []
    for _, row in df_main.iterrows():
        items = []
        for col in encom_individual_cols:
            val = str(row[col]).strip().upper() if pd.notna(row[col]) else ""
            if val == "SI":
                items.append(col)
        combined_encom.append(" | ".join(items))
    df_main["ENCOMENDACIONES"] = combined_encom

    # Eliminar las columnas individuales de encomendaciones
    df_main = df_main.drop(columns=encom_individual_cols, errors="ignore")

    # Convertir numéricas
    for col in ["Horas por contrato", "Horas Totales", "Total Descuentos semanal (horas)", "Total Horas Clínicas"]:
        if col in df_main.columns:
            df_main[col] = pd.to_numeric(df_main[col], errors="coerce")

    # ── Horas Indirectas ──
    df_horas = pd.read_excel(path, sheet_name=SHEET_HORAS, header=0, dtype=str)
    df_horas.columns = [str(c).strip() for c in df_horas.columns]
    for col in df_horas.columns[1:]:
        df_horas[col] = pd.to_numeric(df_horas[col], errors="coerce")

    # ── DOT IDEAL ──
    df_dot = pd.read_excel(path, sheet_name=SHEET_DOT, header=None, dtype=str)

    return df_main, df_horas, df_dot


def parse_dot_ideal(df_dot_raw: pd.DataFrame) -> pd.DataFrame:
    """Extrae la segunda tabla de DOT IDEAL (población 31-01-2026)."""
    # Buscar la fila que contiene "31-01-2026" o la segunda ocurrencia de "CESFAM"
    cesfam_rows = []
    for i, row in df_dot_raw.iterrows():
        vals = [str(v).strip() for v in row.values if pd.notna(v)]
        if any("CESFAM" in v and "SECTORES" in str(df_dot_raw.iloc[i].values) for v in vals):
            cesfam_rows.append(i)

    # Usar la segunda tabla (más reciente) si existe, si no la primera
    if len(cesfam_rows) >= 2:
        header_row = cesfam_rows[1]
    elif cesfam_rows:
        header_row = cesfam_rows[0]
    else:
        return pd.DataFrame()

    headers = df_dot_raw.iloc[header_row].values
    # Limpiar headers
    clean_headers = []
    for h in headers:
        h_str = str(h).strip() if pd.notna(h) else ""
        clean_headers.append(h_str)

    # Buscar datos después del header
    data_rows = []
    for i in range(header_row + 1, len(df_dot_raw)):
        row = df_dot_raw.iloc[i]
        # Si la fila está vacía, parar
        vals = [v for v in row.values if pd.notna(v) and str(v).strip()]
        if len(vals) < 3:
            break
        data_rows.append(row.values)

    if not data_rows:
        return pd.DataFrame()

    df = pd.DataFrame(data_rows, columns=clean_headers)
    # Eliminar columnas vacías
    df = df.loc[:, (df.columns != "") & (df.columns != "None")]
    # Desduplicar columnas
    seen = {}
    new_cols = []
    for c in df.columns:
        if c in seen:
            seen[c] += 1
            new_cols.append(f"{c}_{seen[c]}")
        else:
            seen[c] = 0
            new_cols.append(c)
    df.columns = new_cols

    # Convertir numéricos
    for col in df.columns:
        if col not in ["CESFAM", "CECOSF"]:
            try:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            except (TypeError, ValueError):
                pass

    return df


def build_horas_lookup(df_horas: pd.DataFrame) -> dict:
    """Construye diccionario Item → Horas a la semana desde Horas Indirectas."""
    lookup = {}
    col_item = df_horas.columns[0]
    col_horas = None
    for c in df_horas.columns:
        if "horas a la semana" in c.lower():
            col_horas = c
            break
    if col_horas is None and len(df_horas.columns) >= 6:
        col_horas = df_horas.columns[5]
    if col_horas is None:
        return lookup

    for _, row in df_horas.iterrows():
        item = str(row[col_item]).strip() if pd.notna(row[col_item]) else ""
        horas = row[col_horas] if pd.notna(row[col_horas]) else 0
        if item and isinstance(horas, (int, float)) and horas > 0:
            lookup[item] = float(horas)
    return lookup


def save_main_sheet(path: str, df: pd.DataFrame):
    """Guarda cambios de la hoja principal al Excel preservando otras hojas."""
    wb = openpyxl.load_workbook(path)
    ws = wb[SHEET_MAIN]

    # Determinar fila de encabezado (fila 2 en Excel = row index 2)
    header_row = 2
    # Mapear columnas
    headers = {}
    for col_idx in range(1, ws.max_column + 1):
        val = ws.cell(row=header_row, column=col_idx).value
        if val:
            headers[str(val).strip()] = col_idx

    # Si REVISADO no existe en Excel, insertar columna
    if "REVISADO" not in headers:
        ws.insert_cols(1)
        ws.cell(row=header_row, column=1, value="REVISADO")
        # Recalcular headers (todo corrido +1)
        headers = {"REVISADO": 1}
        for col_idx in range(2, ws.max_column + 1):
            val = ws.cell(row=header_row, column=col_idx).value
            if val:
                headers[str(val).strip()] = col_idx

    # Escribir datos
    data_start = header_row + 1
    cols_to_write = ["REVISADO", "ENCOMENDACIONES", "OBSERVACIONES",
                     "Total Descuentos semanal (horas)", "Total Horas Clínicas"]

    # También localizar las columnas individuales de encomendaciones para escribir SI/vacío
    encom_col_indices = {}  # nombre_actividad → col_idx en Excel
    for col_idx in range(1, ws.max_column + 1):
        val = ws.cell(row=header_row, column=col_idx).value
        if val and str(val).strip() not in _KNOWN_MAIN_COLS:
            encom_col_indices[str(val).strip()] = col_idx

    for df_idx, row in df.iterrows():
        excel_row = data_start + df_idx
        for col_name in cols_to_write:
            if col_name in headers and col_name in df.columns:
                val = row[col_name]
                if pd.isna(val):
                    val = None
                ws.cell(row=excel_row, column=headers[col_name], value=val)

        # Escribir marcas SI/vacío en columnas individuales de encomendaciones
        encom_str = str(row.get("ENCOMENDACIONES", "")) if pd.notna(row.get("ENCOMENDACIONES")) else ""
        active_items = {x.strip() for x in encom_str.split("|") if x.strip()}
        for act_name, col_idx in encom_col_indices.items():
            ws.cell(row=excel_row, column=col_idx, value="SI" if act_name in active_items else None)

    wb.save(path)
    wb.close()


def save_horas_sheet(path: str, df_horas: pd.DataFrame):
    """Guarda cambios de Horas Indirectas al Excel."""
    wb = openpyxl.load_workbook(path)
    ws = wb[SHEET_HORAS]
    for i, row in df_horas.iterrows():
        for j, val in enumerate(row.values):
            cell_val = val if pd.notna(val) else None
            ws.cell(row=i + 2, column=j + 1, value=cell_val)
    wb.save(path)
    wb.close()


def save_dot_ideal_sheet(path: str, df_dot_raw: pd.DataFrame):
    """Guarda DOT IDEAL crudo al Excel."""
    wb = openpyxl.load_workbook(path)
    ws = wb[SHEET_DOT]
    for i, row in df_dot_raw.iterrows():
        for j, val in enumerate(row.values):
            cell_val = val if pd.notna(val) else None
            ws.cell(row=i + 1, column=j + 1, value=cell_val)
    wb.save(path)
    wb.close()


# ─────────────────────────────────────────────────────────────
# DB PERSISTENCIA  (Turso HTTP API / SQLite local fallback)
# ─────────────────────────────────────────────────────────────
_TABLE_MAIN = "dotacion_main"
_TABLE_HORAS = "horas_indirectas"
_TABLE_DOT = "dot_ideal"

# ── Turso HTTP helpers ──────────────────────────────────────
def _turso_cfg():
    """Lee URL y token de secrets (vacío si no configurado)."""
    try:
        return st.secrets.get("TURSO_URL", ""), st.secrets.get("TURSO_TOKEN", "")
    except Exception:
        return "", ""


def _use_turso():
    url, tok = _turso_cfg()
    return bool(url and tok)


def _turso_http(stmts):
    """Ejecuta lista de sentencias SQL via Turso HTTP pipeline API.
    stmts: list[str | dict(sql, args)]  →  retorna lista de results.
    """
    url, tok = _turso_cfg()
    endpoint = url.replace("libsql://", "https://").rstrip("/") + "/v2/pipeline"
    reqs = []
    for s in stmts:
        if isinstance(s, str):
            reqs.append({"type": "execute", "stmt": {"sql": s}})
        else:
            reqs.append({"type": "execute", "stmt": s})
    reqs.append({"type": "close"})
    r = _req.post(endpoint, json={"requests": reqs},
                  headers={"Authorization": f"Bearer {tok}"}, timeout=60)
    r.raise_for_status()
    out = []
    for item in r.json().get("results", []):
        if item.get("type") == "ok":
            resp = item.get("response", {})
            if resp.get("type") == "execute":
                out.append(resp.get("result", {}))
    return out


def _turso_val(v):
    t = v.get("type", "null")
    if t == "null":
        return None
    if t == "integer":
        return int(v["value"])
    if t == "float":
        return float(v["value"])
    return v.get("value")


def _turso_query_df(sql):
    """SELECT → DataFrame via HTTP."""
    results = _turso_http([sql])
    if not results:
        return pd.DataFrame()
    res = results[0]
    cols = [c["name"] for c in res.get("cols", [])]
    rows = [[_turso_val(c) for c in row] for row in res.get("rows", [])]
    return pd.DataFrame(rows, columns=cols) if cols else pd.DataFrame()


def _turso_exec_single(sql):
    """Ejecuta una sentencia y retorna el primer fila[0] si hay resultado."""
    results = _turso_http([sql])
    if results and results[0].get("rows"):
        row0 = results[0]["rows"][0]
        return _turso_val(row0[0])
    return None


def _to_turso_arg(v):
    try:
        if v is None or pd.isna(v):
            return {"type": "null"}
    except (ValueError, TypeError):
        pass
    return {"type": "text", "value": str(v)}


def _turso_save_df(df, table_name):
    """Guarda DataFrame completo en Turso via HTTP (DROP+CREATE+INSERT).
    Usa tabla temporal para evitar pérdida de datos si falla el INSERT.
    """
    cols = df.columns.tolist()
    col_defs = ", ".join([f'[{c}] TEXT' for c in cols])
    tmp_name = f"_tmp_{table_name}"
    # Crear tabla temporal
    _turso_http([
        f'DROP TABLE IF EXISTS [{tmp_name}]',
        f'CREATE TABLE [{tmp_name}] ({col_defs})',
    ])
    col_names = ", ".join([f'[{c}]' for c in cols])
    placeholders = ", ".join(["?" for _ in cols])
    BATCH = 20
    try:
        for start in range(0, len(df), BATCH):
            stmts = []
            for _, row in df.iloc[start:start + BATCH].iterrows():
                args = [_to_turso_arg(v) for v in row]
                stmts.append({"sql": f"INSERT INTO [{tmp_name}] ({col_names}) VALUES ({placeholders})", "args": args})
            _turso_http(stmts)
        # Todos los INSERTs OK → reemplazar tabla real
        _turso_http([
            f'DROP TABLE IF EXISTS [{table_name}]',
            f'ALTER TABLE [{tmp_name}] RENAME TO [{table_name}]',
        ])
    except Exception:
        # Si falla, limpiar temporal pero no tocar la tabla real
        try:
            _turso_http([f'DROP TABLE IF EXISTS [{tmp_name}]'])
        except Exception:
            pass
        raise


# ── Funciones de persistencia (dual: Turso + SQLite local) ──

def db_has_data() -> bool:
    """Retorna True si la BD tiene registros (Turso con fallback a SQLite local)."""
    # 1) Intentar Turso si está configurado
    if _use_turso():
        try:
            n = _turso_exec_single(f"SELECT COUNT(*) FROM {_TABLE_MAIN}")
            if (n or 0) > 0:
                return True
            # Turso vacío → verificar SQLite local como fallback
        except Exception:
            pass  # Turso inaccesible → continuar con SQLite
    # 2) SQLite local (siempre como fuente de verdad de respaldo)
    try:
        if not _DB_PATH.exists():
            return False
        con = sqlite3.connect(str(_DB_PATH))
        cur = con.execute(
            f"SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='{_TABLE_MAIN}'"
        )
        if cur.fetchone()[0] == 0:
            con.close()
            return False
        cur = con.execute(f"SELECT COUNT(*) FROM {_TABLE_MAIN}")
        n = cur.fetchone()[0]
        con.close()
        return n > 0
    except Exception:
        return False


def db_save_main(df: pd.DataFrame):
    """Guarda el DataFrame principal (reemplaza la tabla completa)."""
    if _use_turso():
        try:
            _turso_save_df(df, _TABLE_MAIN)
        except Exception as e:
            import traceback; traceback.print_exc()
    con = sqlite3.connect(str(_DB_PATH))
    df.to_sql(_TABLE_MAIN, con, if_exists="replace", index=False)
    con.close()


def _normalize_cesfam_col(df: pd.DataFrame) -> pd.DataFrame:
    """Normaliza la columna CESFAM: mayúsculas + unifica Nº→N°."""
    if "CESFAM" in df.columns:
        df["CESFAM"] = (
            df["CESFAM"]
            .fillna("")
            .str.strip()
            .str.upper()
            .str.replace("Nº", "N°", regex=False)
            .str.replace("º", "°", regex=False)
        )
    return df


def _sqlite_load_table(table: str) -> pd.DataFrame:
    """Lee una tabla desde SQLite local. Retorna DataFrame vacío si falla."""
    try:
        con = sqlite3.connect(str(_DB_PATH))
        df = pd.read_sql(f"SELECT * FROM [{table}]", con)
        con.close()
        return df
    except Exception:
        return pd.DataFrame()


def db_load_main() -> pd.DataFrame:
    """Carga el DataFrame principal (Turso con fallback a SQLite local)."""
    df = pd.DataFrame()
    if _use_turso():
        try:
            df = _turso_query_df(f"SELECT * FROM {_TABLE_MAIN}")
        except Exception:
            pass
    if df.empty:
        df = _sqlite_load_table(_TABLE_MAIN)
    # Restaurar tipos numéricos
    for col in ["Horas por contrato", "Horas Totales",
                "Total Descuentos semanal (horas)", "Total Horas Clínicas"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    # Asegurar columnas de control
    if "_ES_NUEVO" not in df.columns:
        df["_ES_NUEVO"] = False
    else:
        df["_ES_NUEVO"] = df["_ES_NUEVO"].fillna(False).astype(bool)
    if "REVISADO" in df.columns:
        df["REVISADO"] = df["REVISADO"].fillna("")
    # Normalizar CESFAM
    df = _normalize_cesfam_col(df)
    return df


def db_merge_new_ruts(df_existing: pd.DataFrame, df_excel: pd.DataFrame) -> pd.DataFrame:
    """Compara RUTs del Excel con los existentes en BD.
    Solo agrega personas NUEVAS (RUT no presente en BD), marcándolas como _ES_NUEVO.
    """
    def norm_rut(val):
        return str(val).replace(".", "").replace("-", "").strip().upper()
    existing_ruts = set(df_existing["RUT"].dropna().apply(norm_rut))
    excel_ruts = set(df_excel["RUT"].dropna().apply(norm_rut))

    new_ruts = excel_ruts - existing_ruts
    if not new_ruts:
        return df_existing, 0

    # Filtrar solo filas con RUTs nuevos
    df_new = df_excel[df_excel["RUT"].dropna().apply(norm_rut).isin(new_ruts)].copy()
    df_new["_ES_NUEVO"] = True
    df_new["REVISADO"] = ""

    # Asegurar mismas columnas
    for col in df_existing.columns:
        if col not in df_new.columns:
            df_new[col] = "" if df_existing[col].dtype == object else 0
    df_new = df_new[[c for c in df_existing.columns if c in df_new.columns]]

    # Concatenar
    df_merged = pd.concat([df_existing, df_new], ignore_index=True)
    return df_merged, len(df_new)


def db_save_horas(df: pd.DataFrame):
    """Guarda la tabla Horas Indirectas."""
    if _use_turso():
        try:
            _turso_save_df(df, _TABLE_HORAS)
        except Exception as e:
            import traceback; traceback.print_exc()
    con = sqlite3.connect(str(_DB_PATH))
    df.to_sql(_TABLE_HORAS, con, if_exists="replace", index=False)
    con.close()


def db_load_horas() -> pd.DataFrame | None:
    """Carga Horas Indirectas (Turso con fallback a SQLite local)."""
    df = pd.DataFrame()
    if _use_turso():
        try:
            df = _turso_query_df(f"SELECT * FROM {_TABLE_HORAS}")
        except Exception:
            pass
    if df.empty:
        df = _sqlite_load_table(_TABLE_HORAS)
    if df.empty:
        return None
    for col in df.columns[1:]:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


def db_save_dot(df: pd.DataFrame):
    """Guarda la tabla DOT IDEAL procesada."""
    if _use_turso():
        try:
            _turso_save_df(df, _TABLE_DOT)
        except Exception as e:
            import traceback; traceback.print_exc()
    con = sqlite3.connect(str(_DB_PATH))
    df.to_sql(_TABLE_DOT, con, if_exists="replace", index=False)
    con.close()


def db_load_dot() -> pd.DataFrame | None:
    """Carga DOT IDEAL (Turso con fallback a SQLite local)."""
    df = pd.DataFrame()
    if _use_turso():
        try:
            df = _turso_query_df(f"SELECT * FROM {_TABLE_DOT}")
        except Exception:
            pass
    if df.empty:
        df = _sqlite_load_table(_TABLE_DOT)
    if df.empty:
        return None
    for col in df.columns:
        if col not in ["CESFAM", "CECOSF", "SECTORES"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


def db_export_excel(df: pd.DataFrame) -> bytes:
    """Exporta el DataFrame actual como archivo Excel descargable."""
    output = io.BytesIO()
    # Excluir columnas internas
    export_cols = [c for c in df.columns if not c.startswith("_")]
    df[export_cols].to_excel(output, index=False, sheet_name="Revisión Dotación")
    output.seek(0)
    return output.getvalue()


# ─────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────
def init_state():
    defaults = {
        "excel_path": None,
        "df_main": None,
        "df_horas": None,
        "df_dot_raw": None,
        "df_dot": None,
        "horas_lookup": {},
        "dirty_main": False,
        "dirty_horas": False,
        "last_save_msg": "",
        "_loaded_from": None,
        # Filtros persistentes (sobreviven cambio de página)
        "_p_ef_tipo": [],
        "_p_ef_cesfam": [],
        "_p_ef_planta": [],
        "_p_ef_cargo": [],
        "_p_ef_unidad": [],
        "_p_ef_search": "",
        "_p_ef_rev_filter": "Todos",
        "_p_ef_sort_col": "(sin ordenar)",
        "_p_ef_sort_dir": "↑ Ascendente",
        "_ag_version": 0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_state()


# ─────────────────────────────────────────────────────────────
# CARGA DE DATOS
# ─────────────────────────────────────────────────────────────
def load_data():
    """Carga datos: prefiere SQLite si existe, si no carga desde Excel y persiste."""
    # 1) Intentar cargar desde SQLite (revisión guardada)
    if db_has_data():
        df_main = db_load_main()
        st.session_state.df_main = df_main.reset_index(drop=True)
        st.session_state._loaded_from = "sqlite"
        # Cargar horas y dot desde SQLite
        _df_h = db_load_horas()
        _df_d = db_load_dot()
        if _df_h is not None:
            st.session_state.df_horas = _df_h
            st.session_state.horas_lookup = build_horas_lookup(_df_h)
        if _df_d is not None:
            st.session_state.df_dot = _df_d
        if _df_h is None or _df_d is None:
            _load_secondary_sheets()
        return True

    # 2) Sin BD → intentar cargar desde Excel si hay uno disponible
    path = st.session_state.get("excel_path")
    if path and Path(path).exists():
        try:
            with st.spinner("Cargando datos desde Excel..."):
                readable = _ensure_readable(Path(path))
                mtime = Path(readable).stat().st_mtime
                sheet = st.session_state.get("_selected_sheet", SHEET_MAIN)
                df_main, df_horas, df_dot_raw = load_excel(readable, mtime, sheet_main=sheet)
                df_main = _normalize_cesfam_col(df_main.reset_index(drop=True))
                df_main["_ES_NUEVO"] = False
                st.session_state.df_main = df_main
                st.session_state._loaded_from = "excel"
                st.session_state.df_horas = df_horas.reset_index(drop=True)
                st.session_state.df_dot_raw = df_dot_raw
                st.session_state.df_dot = parse_dot_ideal(df_dot_raw)
                st.session_state.horas_lookup = build_horas_lookup(df_horas)
                # Persistir en SQLite para carga rápida en el futuro
                db_save_main(st.session_state.df_main)
                db_save_horas(st.session_state.df_horas)
                if st.session_state.df_dot is not None and not st.session_state.df_dot.empty:
                    db_save_dot(st.session_state.df_dot)
            st.toast("Datos cargados y guardados correctamente", icon="✅")
            return True
        except Exception as e:
            st.error(f"Error al cargar Excel: {e}")
            # No cambiar df_main para que pueda reintentarse en el próximo rerun
            return False

    # 3) Sin BD ni Excel disponible
    st.warning("No hay datos guardados. Por favor, cargue un archivo Excel para iniciar.")
    # No asignar df_main aquí para permitir reintentos automáticos
    return False


def _load_secondary_sheets():
    """Carga Horas Indirectas y DOT IDEAL desde Excel y persiste en SQLite."""
    path = st.session_state.excel_path
    if path is None:
        # Intentar ubicación conocida del archivo como fallback
        _found = find_excel_file()
        if _found:
            path = str(_found)
        else:
            return
    p = Path(path)
    if not p.exists():
        return
    readable_path = _ensure_readable(p)
    mtime = Path(readable_path).stat().st_mtime
    _, df_horas, df_dot_raw = load_excel(readable_path, mtime)
    st.session_state.df_horas = df_horas.reset_index(drop=True)
    st.session_state.df_dot_raw = df_dot_raw
    st.session_state.df_dot = parse_dot_ideal(df_dot_raw)
    st.session_state.horas_lookup = build_horas_lookup(df_horas)
    # Persistir en SQLite
    if st.session_state.df_horas is not None:
        db_save_horas(st.session_state.df_horas)
    if st.session_state.df_dot is not None and not st.session_state.df_dot.empty:
        db_save_dot(st.session_state.df_dot)


# ─────────────────────────────────────────────────────────────
# KPIs
# ─────────────────────────────────────────────────────────────
def compute_kpis(df: pd.DataFrame) -> dict:
    total = len(df)
    revisados = df["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"]).sum() if "REVISADO" in df.columns else 0
    pendientes = total - revisados
    pct = (revisados / total * 100) if total > 0 else 0

    # Horas
    horas_totales = df["Horas Totales"].sum() if "Horas Totales" in df.columns else 0
    descuentos = df["Total Descuentos semanal (horas)"].sum() if "Total Descuentos semanal (horas)" in df.columns else 0
    horas_clinicas = df["Total Horas Clínicas"].sum() if "Total Horas Clínicas" in df.columns else 0
    jornadas = horas_clinicas / 44 if horas_clinicas > 0 else 0

    # Unidades
    n_unidades = df["UNIDAD DE DESEMPEÑO"].nunique() if "UNIDAD DE DESEMPEÑO" in df.columns else 0

    # CESFAM
    n_cesfam = df["CESFAM"].nunique() if "CESFAM" in df.columns else 0

    return {
        "total": total,
        "revisados": int(revisados),
        "pendientes": int(pendientes),
        "pct_revisados": pct,
        "horas_totales": horas_totales,
        "descuentos": descuentos,
        "horas_clinicas": horas_clinicas,
        "jornadas": jornadas,
        "n_unidades": n_unidades,
        "n_cesfam": n_cesfam,
    }


# ─────────────────────────────────────────────────────────────
# SIDEBAR + FILTROS
# ─────────────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown("### 🏥 SSMC · Dotación APS")
        st.caption("Gestión de Dotación Marzo 2026")
        st.divider()

        # Evaluar una sola vez si hay datos en BD (evita múltiples conexiones SQLite)
        _bd_tiene_datos = db_has_data()

        # ── Carga de archivo Excel mediante botón ──
        uploaded = st.file_uploader(
            "📁 Seleccionar archivo Excel",
            type=["xlsx"],
            key="_file_uploader",
            help="Haga clic en «Browse files» para elegir el archivo de dotación",
        )

        # Si se sube un archivo nuevo, guardarlo en temp y registrarlo
        if uploaded is not None:
            _up_name = uploaded.name
            _prev_up = st.session_state.get("_uploaded_file_name")
            _tmp_path = _WORK_DIR / _up_name
            # Escribir solo si es archivo nuevo o diferente
            if _prev_up != _up_name or not _tmp_path.exists():
                _tmp_path.write_bytes(uploaded.getvalue())
                st.session_state["_uploaded_file_name"] = _up_name
            st.session_state.excel_path = str(_tmp_path)
        elif st.session_state.get("_uploaded_file_name"):
            # Mantener el archivo previamente subido
            _tmp_path = _WORK_DIR / st.session_state["_uploaded_file_name"]
            if _tmp_path.exists():
                st.session_state.excel_path = str(_tmp_path)

        _has_file = st.session_state.get("excel_path") and Path(st.session_state.excel_path).exists()

        if not _has_file:
            if not _bd_tiene_datos:
                st.info("Suba un archivo .xlsx para comenzar")
                st.stop()
            # Sin archivo pero con BD existente → cargar desde SQLite
        else:
            full_path = st.session_state.excel_path
            _prev_file = st.session_state.get("_selected_file")
            _cur_file_name = Path(full_path).name
            st.session_state["_selected_file"] = _cur_file_name

            # ── Selector de hoja ──
            sheets = _list_excel_sheets(full_path)
            if sheets:
                cur_sheet = st.session_state.get("_selected_sheet", SHEET_MAIN)
                default_sheet_idx = 0
                if cur_sheet in sheets:
                    default_sheet_idx = sheets.index(cur_sheet)
                sel_sheet = st.selectbox(
                    "📄 Hoja de dotación",
                    sheets,
                    index=default_sheet_idx,
                    key="_sheet_selector",
                    help="Seleccione la hoja que contiene los datos de dotación",
                )
                _prev_sheet = st.session_state.get("_selected_sheet")
                st.session_state["_selected_sheet"] = sel_sheet

                # Si cambió archivo o hoja → mezclar nuevos RUTs (NUNCA borrar BD)
                _file_changed = (_prev_file is not None and _prev_file != _cur_file_name)
                _sheet_changed = (_prev_sheet is not None and _prev_sheet != sel_sheet)
                if _file_changed or _sheet_changed:
                    load_excel.clear()
                    try:
                        readable = _ensure_readable(Path(full_path))
                        mtime = Path(readable).stat().st_mtime
                        df_new, _, _ = load_excel(readable, mtime, sheet_main=sel_sheet)
                        df_new = _normalize_cesfam_col(df_new.reset_index(drop=True))
                        df_new["_ES_NUEVO"] = False
                        if st.session_state.df_main is not None and len(st.session_state.df_main) > 0:
                            df_merged, n_new = db_merge_new_ruts(st.session_state.df_main, df_new)
                            st.session_state.df_main = df_merged.reset_index(drop=True)
                        elif _bd_tiene_datos:
                            existing = db_load_main()
                            df_merged, n_new = db_merge_new_ruts(existing, df_new)
                            st.session_state.df_main = df_merged.reset_index(drop=True)
                        else:
                            st.session_state.df_main = df_new
                            n_new = len(df_new)
                        db_save_main(st.session_state.df_main)
                        # Recargar hojas secundarias del nuevo archivo
                        _load_secondary_sheets()
                        if n_new > 0:
                            st.toast(f"✅ {n_new} nuevos funcionarios incorporados", icon="🆕")
                        else:
                            st.toast("Sin nuevos RUTs. Base de datos intacta.", icon="ℹ️")
                    except Exception as e:
                        st.error(f"Error al cargar: {e}")
                    st.rerun()

        if st.session_state.df_main is None:
            if not load_data():
                st.stop()

        st.divider()
        nav = st.radio(
            "Navegación",
            ["📊 Panel Principal", "📋 Dotación (Editor)", "⏱️ Horas Indirectas",
             "🎯 Dotación Ideal", "📈 Análisis de Brechas", "📝 Acta de Reunión"],
            label_visibility="collapsed",
        )

        st.divider()

        # Filtros
        df = st.session_state.df_main
        filtros = {}

        cesfam_opts = sorted(df["CESFAM"].dropna().unique().tolist()) if "CESFAM" in df.columns else []
        if cesfam_opts:
            st.markdown("**Filtros**")
            f_cesfam = st.multiselect("CESFAM", cesfam_opts, key="f_cesfam")
            if f_cesfam:
                filtros["CESFAM"] = f_cesfam

        planta_opts = sorted(df["DESCRIPCIÓN DE PLANTA"].dropna().unique().tolist()) if "DESCRIPCIÓN DE PLANTA" in df.columns else []
        if planta_opts:
            f_planta = st.multiselect("Tipo Planta", planta_opts, key="f_planta")
            if f_planta:
                filtros["DESCRIPCIÓN DE PLANTA"] = f_planta

        cargo_opts = sorted(df["CARGO"].dropna().unique().tolist()) if "CARGO" in df.columns else []
        if cargo_opts:
            f_cargo = st.multiselect("Cargo", cargo_opts, key="f_cargo")
            if f_cargo:
                filtros["CARGO"] = f_cargo

        rev_opts = ["Todos", "Revisados", "Pendientes"]
        f_rev = st.radio("Estado Revisión", rev_opts, horizontal=True, key="f_rev")
        if f_rev == "Revisados":
            filtros["_revisados"] = True
        elif f_rev == "Pendientes":
            filtros["_pendientes"] = True

        st.divider()

        # Estado de la BD
        if _bd_tiene_datos:
            st.caption("💾 Auto-guardado activo · BD local")
        else:
            st.caption("📭 Sin revisión guardada")

        # Botón de guardado manual (respaldo)
        if st.button("💾 Guardar ahora", use_container_width=True,
                      help="Forzar guardado manual (el auto-guardado ya protege tus cambios)"):
            try:
                db_save_main(st.session_state.df_main)
                st.session_state.dirty_main = False
                if st.session_state.dirty_horas:
                    save_horas_sheet(st.session_state.excel_path, st.session_state.df_horas)
                    st.session_state.dirty_horas = False
                load_excel.clear()
                st.toast("Revisión guardada manualmente", icon="💾")
                st.rerun()
            except Exception as e:
                st.error(f"Error al guardar: {e}")

        # Recargar desde Excel (solo trae RUTs nuevos)
        if st.button("📥 Importar nuevos desde Excel", use_container_width=True,
                      help="Compara el Excel de origen con la BD. Solo agrega personas con RUT nuevo."):
            if st.session_state.excel_path:
                try:
                    readable_path = _ensure_readable(Path(st.session_state.excel_path))
                    mtime = Path(readable_path).stat().st_mtime
                    load_excel.clear()
                    df_excel_fresh, _, _ = load_excel(readable_path, mtime)
                    df_excel_fresh = df_excel_fresh.reset_index(drop=True)
                    df_excel_fresh["_ES_NUEVO"] = False  # se marcará True solo los nuevos
                    df_merged, n_new = db_merge_new_ruts(st.session_state.df_main, df_excel_fresh)
                    if n_new > 0:
                        st.session_state.df_main = df_merged.reset_index(drop=True)
                        db_save_main(st.session_state.df_main)
                        st.toast(f"✅ {n_new} personas nuevas importadas", icon="🆕")
                    else:
                        st.toast("No se encontraron personas nuevas en el Excel", icon="ℹ️")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al importar: {e}")

        # Descargar revisión como Excel
        if st.session_state.df_main is not None:
            excel_bytes = db_export_excel(st.session_state.df_main)
            st.download_button(
                "📤 Descargar revisión (Excel)",
                data=excel_bytes,
                file_name="revision_dotacion.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        # Reiniciar BD (cargar todo desde Excel de nuevo)
        if st.button("🔄 Reiniciar desde Excel (full)", use_container_width=True,
                      help="Borra la BD y vuelve a cargar todo desde el Excel original"):
            if _DB_PATH.exists():
                _DB_PATH.unlink()
            load_excel.clear()
            st.session_state.df_main = None
            st.session_state.dirty_main = False
            st.session_state.dirty_horas = False
            st.rerun()

        return nav, filtros


def apply_main_filters(df: pd.DataFrame, filtros: dict) -> pd.DataFrame:
    dff = df.copy()
    for col, vals in filtros.items():
        if col.startswith("_"):
            continue
        if col in dff.columns:
            dff = dff[dff[col].isin(vals)]
    if filtros.get("_revisados"):
        dff = dff[dff["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"])]
    elif filtros.get("_pendientes"):
        dff = dff[~dff["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"])]
    return dff


# ─────────────────────────────────────────────────────────────
# PÁGINA: PANEL PRINCIPAL
# ─────────────────────────────────────────────────────────────
def page_panel(filtros):
    st.markdown("""
    <div class="main-header">
        <h1>📊 Panel de Control — Dotación Marzo 2026</h1>
        <p>Servicio de Salud Metropolitano Central · Revisión de Dotación APS</p>
    </div>
    """, unsafe_allow_html=True)

    df = st.session_state.df_main
    dff = apply_main_filters(df, filtros)
    kpis = compute_kpis(dff)

    # ── KPIs principales ──
    st.markdown('<div class="kpi-row">', unsafe_allow_html=True)
    cols = st.columns(6)
    with cols[0]:
        pct_class = "green" if kpis["pct_revisados"] >= 80 else ("amber" if kpis["pct_revisados"] >= 40 else "red")
        st.markdown(f'''<div class="kpi-card {pct_class}">
            <div class="kpi-value">{kpis["pct_revisados"]:.1f}%</div>
            <div class="kpi-label">Avance Revisión</div>
            <div class="kpi-sub">{kpis["revisados"]} de {kpis["total"]}</div>
        </div>''', unsafe_allow_html=True)
    with cols[1]:
        st.markdown(f'''<div class="kpi-card green">
            <div class="kpi-value">{kpis["revisados"]:,}</div>
            <div class="kpi-label">Revisados</div>
        </div>''', unsafe_allow_html=True)
    with cols[2]:
        st.markdown(f'''<div class="kpi-card red">
            <div class="kpi-value">{kpis["pendientes"]:,}</div>
            <div class="kpi-label">Pendientes</div>
        </div>''', unsafe_allow_html=True)
    with cols[3]:
        st.markdown(f'''<div class="kpi-card">
            <div class="kpi-value">{kpis["horas_clinicas"]:,.0f}</div>
            <div class="kpi-label">Hrs Clínicas</div>
            <div class="kpi-sub">{kpis["jornadas"]:.1f} jornadas</div>
        </div>''', unsafe_allow_html=True)
    with cols[4]:
        st.markdown(f'''<div class="kpi-card">
            <div class="kpi-value">{kpis["n_unidades"]}</div>
            <div class="kpi-label">Unidades</div>
        </div>''', unsafe_allow_html=True)
    with cols[5]:
        st.markdown(f'''<div class="kpi-card">
            <div class="kpi-value">{kpis["n_cesfam"]}</div>
            <div class="kpi-label">CESFAM</div>
        </div>''', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Progreso por categorías ──
    c1, c2 = st.columns(2)

    with c1:
        st.markdown('<div class="section-title">Revisión por Descripción de Planta</div>', unsafe_allow_html=True)
        if "DESCRIPCIÓN DE PLANTA" in dff.columns:
            rev_mask = dff["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"])
            grp = dff.groupby("DESCRIPCIÓN DE PLANTA").agg(
                Total=("REVISADO", "size"),
                Revisados=("REVISADO", lambda x: x.str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"]).sum())
            ).reset_index()
            grp["Pendientes"] = grp["Total"] - grp["Revisados"]
            grp = grp.sort_values("Total", ascending=True)
            fig = go.Figure()
            fig.add_trace(go.Bar(y=grp["DESCRIPCIÓN DE PLANTA"], x=grp["Revisados"],
                                 name="Revisados", orientation="h",
                                 marker_color=C_SUCCESS, text=grp["Revisados"], textposition="inside"))
            fig.add_trace(go.Bar(y=grp["DESCRIPCIÓN DE PLANTA"], x=grp["Pendientes"],
                                 name="Pendientes", orientation="h",
                                 marker_color="#D1D5DB", text=grp["Pendientes"], textposition="inside"))
            fig.update_layout(barmode="stack", height=300, margin=dict(l=0,r=20,t=10,b=0),
                              legend=dict(orientation="h", yanchor="bottom", y=1.02),
                              font=dict(family="Inter", size=11))
            st.plotly_chart(fig, use_container_width=True)

    with c2:
        st.markdown('<div class="section-title">Revisión por Cargo (Top 12)</div>', unsafe_allow_html=True)
        if "CARGO" in dff.columns:
            grp2 = dff.groupby("CARGO").agg(
                Total=("REVISADO", "size"),
                Revisados=("REVISADO", lambda x: x.str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"]).sum())
            ).reset_index()
            grp2["Pendientes"] = grp2["Total"] - grp2["Revisados"]
            grp2 = grp2.nlargest(12, "Total").sort_values("Total", ascending=True)
            fig2 = go.Figure()
            fig2.add_trace(go.Bar(y=grp2["CARGO"], x=grp2["Revisados"],
                                  name="Revisados", orientation="h",
                                  marker_color=C_SUCCESS, text=grp2["Revisados"], textposition="inside"))
            fig2.add_trace(go.Bar(y=grp2["CARGO"], x=grp2["Pendientes"],
                                  name="Pendientes", orientation="h",
                                  marker_color="#D1D5DB", text=grp2["Pendientes"], textposition="inside"))
            fig2.update_layout(barmode="stack", height=300, margin=dict(l=0,r=20,t=10,b=0),
                               legend=dict(orientation="h", yanchor="bottom", y=1.02),
                               font=dict(family="Inter", size=11))
            st.plotly_chart(fig2, use_container_width=True)

    # ── Distribución por CESFAM ──
    st.markdown('<div class="section-title">Distribución por CESFAM</div>', unsafe_allow_html=True)
    if "CESFAM" in dff.columns:
        c3, c4 = st.columns(2)
        with c3:
            grp3 = dff.groupby("CESFAM").agg(
                Funcionarios=("RUT", "count"),
                Horas_Clinicas=("Total Horas Clínicas", "sum")
            ).reset_index()
            grp3["Jornadas"] = grp3["Horas_Clinicas"] / 44
            grp3 = grp3.sort_values("Funcionarios", ascending=False)
            fig3 = px.bar(grp3, x="CESFAM", y="Funcionarios", color_discrete_sequence=[C_ACCENT],
                          text="Funcionarios")
            fig3.update_layout(height=300, margin=dict(l=0,r=0,t=10,b=0),
                               xaxis_tickangle=-45, font=dict(family="Inter", size=11))
            fig3.update_traces(textposition="outside")
            st.plotly_chart(fig3, use_container_width=True)

        with c4:
            grp3b = grp3.sort_values("Jornadas", ascending=False)
            fig3b = px.bar(grp3b, x="CESFAM", y="Jornadas", color_discrete_sequence=[C_PRIMARY],
                           text=grp3b["Jornadas"].round(1))
            fig3b.update_layout(height=300, margin=dict(l=0,r=0,t=10,b=0),
                                xaxis_tickangle=-45, font=dict(family="Inter", size=11),
                                yaxis_title="Jornadas (Hrs Clínicas / 44)")
            fig3b.update_traces(textposition="outside")
            st.plotly_chart(fig3b, use_container_width=True)

    # ── Unidades de desempeño ──
    st.markdown('<div class="section-title">Unidades de Desempeño Registradas</div>', unsafe_allow_html=True)
    if "UNIDAD DE DESEMPEÑO" in dff.columns:
        ud_counts = dff["UNIDAD DE DESEMPEÑO"].value_counts().head(20).reset_index()
        ud_counts.columns = ["Unidad", "Funcionarios"]
        fig_ud = px.bar(ud_counts, x="Funcionarios", y="Unidad", orientation="h",
                        color_discrete_sequence=[C_ACCENT], text="Funcionarios")
        fig_ud.update_layout(height=max(300, len(ud_counts) * 28), margin=dict(l=0,r=20,t=10,b=0),
                             font=dict(family="Inter", size=11))
        fig_ud.update_traces(textposition="outside")
        st.plotly_chart(fig_ud, use_container_width=True)


# ─────────────────────────────────────────────────────────────
# PÁGINA: EDITOR PRINCIPAL
# ─────────────────────────────────────────────────────────────
def page_editor(filtros):
    st.markdown("""
    <div class="main-header">
        <h1>📋 Editor de Dotación — TIT-CONTRATA MARZO 2026</h1>
        <p>Edición directa de registros · Los cambios se reflejan en tiempo real</p>
    </div>
    """, unsafe_allow_html=True)

    df = st.session_state.df_main

    # ── Filtros inline estilo Excel ──
    st.markdown('<div class="section-title">Filtros (estilo Excel)</div>', unsafe_allow_html=True)

    # Fila 1: Búsqueda por texto + estado revisión
    fc1, fc2 = st.columns([3, 1])
    with fc1:
        search_text = st.text_input("🔍 Buscar (nombre, RUT, observaciones...)",
                                     key="editor_search", placeholder="Escriba para buscar...")
    with fc2:
        rev_filter = st.selectbox("Estado Revisión", ["Todos", "✅ Revisados", "⏳ Pendientes"],
                                   key="editor_rev_filter")

    # Fila 2: Filtros por columna CONDICIONALES (cascada)
    # Cada filtro muestra solo opciones que existen dado los otros filtros activos
    filter_cols = st.columns(5)

    # Pre-filtrar base con sidebar filters
    _base = apply_main_filters(df.copy(), filtros)

    # Calcular subset parcial para cada filtro (excluyendo ese filtro)
    def _opts(base_df, col, exclude_filters):
        """Retorna opciones de 'col' filtrando por todos los filtros EXCEPTO el propio."""
        tmp = base_df.copy()
        for fc, fv in exclude_filters.items():
            if fv and fc in tmp.columns:
                tmp = tmp[tmp[fc].isin(fv)]
        return sorted(tmp[col].dropna().unique().tolist()) if col in tmp.columns else []

    # Restaurar filtros persistentes si los widgets no existen (cambio de página)
    _persist_map = {
        "ef_tipo": "_p_ef_tipo", "ef_cesfam": "_p_ef_cesfam",
        "ef_planta": "_p_ef_planta", "ef_cargo": "_p_ef_cargo",
        "ef_unidad": "_p_ef_unidad", "editor_search": "_p_ef_search",
        "editor_rev_filter": "_p_ef_rev_filter", "ef_sort_col": "_p_ef_sort_col",
        "ef_sort_dir": "_p_ef_sort_dir",
    }
    for wk, pk in _persist_map.items():
        if wk not in st.session_state and pk in st.session_state:
            st.session_state[wk] = st.session_state[pk]

    # Leer valores actuales de session_state (pueden no existir aún)
    _cur_tipo   = st.session_state.get("ef_tipo", [])
    _cur_cesfam = st.session_state.get("ef_cesfam", [])
    _cur_planta = st.session_state.get("ef_planta", [])
    _cur_cargo  = st.session_state.get("ef_cargo", [])
    _cur_ud     = st.session_state.get("ef_unidad", [])

    # Mapeo nombre_filtro → (columna, valor_actual)
    _all_filters = {
        "TIPO": ("TIPO", _cur_tipo),
        "CESFAM": ("CESFAM", _cur_cesfam),
        "DESCRIPCIÓN DE PLANTA": ("DESCRIPCIÓN DE PLANTA", _cur_planta),
        "CARGO": ("CARGO", _cur_cargo),
        "UNIDAD DE DESEMPEÑO": ("UNIDAD DE DESEMPEÑO", _cur_ud),
    }

    def _excl(skip_col):
        """Filtros activos excluyendo skip_col."""
        return {c: v for c, (_, v) in _all_filters.items() if c != skip_col and v}

    def _opts_keep(base_df, col, exclude_filters, current_sel):
        """Opciones contextuales que siempre conservan la selección actual."""
        available = _opts(base_df, col, exclude_filters)
        # Asegurar que los valores ya seleccionados estén siempre presentes
        for v in current_sel:
            if v not in available:
                available.append(v)
        return sorted(available)

    with filter_cols[0]:
        tipo_opts = _opts_keep(_base, "TIPO", _excl("TIPO"), _cur_tipo)
        f_tipo = st.multiselect("TIPO", tipo_opts, key="ef_tipo", placeholder="Todos")

    with filter_cols[1]:
        cesfam_opts = _opts_keep(_base, "CESFAM", _excl("CESFAM"), _cur_cesfam)
        f_cesfam = st.multiselect("CESFAM", cesfam_opts, key="ef_cesfam", placeholder="Todos")

    with filter_cols[2]:
        planta_opts = _opts_keep(_base, "DESCRIPCIÓN DE PLANTA", _excl("DESCRIPCIÓN DE PLANTA"), _cur_planta)
        f_planta = st.multiselect("Planta", planta_opts, key="ef_planta", placeholder="Todos")

    with filter_cols[3]:
        cargo_opts = _opts_keep(_base, "CARGO", _excl("CARGO"), _cur_cargo)
        f_cargo = st.multiselect("Cargo", cargo_opts, key="ef_cargo", placeholder="Todos")

    with filter_cols[4]:
        ud_opts = _opts_keep(_base, "UNIDAD DE DESEMPEÑO", _excl("UNIDAD DE DESEMPEÑO"), _cur_ud)
        f_ud = st.multiselect("Unidad", ud_opts, key="ef_unidad", placeholder="Todos")

    # Guardar filtros en keys persistentes (sobreviven cambio de página)
    st.session_state["_p_ef_tipo"] = f_tipo
    st.session_state["_p_ef_cesfam"] = f_cesfam
    st.session_state["_p_ef_planta"] = f_planta
    st.session_state["_p_ef_cargo"] = f_cargo
    st.session_state["_p_ef_unidad"] = f_ud

    # Fila 3: Ordenamiento
    sort_c1, sort_c2, sort_c3 = st.columns([2, 2, 1])
    sortable_cols = ["CESFAM", "NOMBRE PROFESIONAL", "CARGO", "DESCRIPCIÓN DE PLANTA",
                     "UNIDAD DE DESEMPEÑO", "RUT", "TIPO", "Horas por contrato",
                     "Horas Totales", "Total Horas Clínicas", "REVISADO"]
    sortable_cols = [c for c in sortable_cols if c in df.columns]
    with sort_c1:
        sort_col = st.selectbox("Ordenar por", ["(sin ordenar)"] + sortable_cols, key="ef_sort_col")
    with sort_c2:
        sort_dir = st.radio("Dirección", ["↑ Ascendente", "↓ Descendente"],
                             horizontal=True, key="ef_sort_dir")
    with sort_c3:
        st.markdown("<br>", unsafe_allow_html=True)
        clear_filters = st.button("🧹 Limpiar filtros", key="ef_clear")

    # Persistir búsqueda, revisión y ordenamiento
    st.session_state["_p_ef_search"] = search_text
    st.session_state["_p_ef_rev_filter"] = rev_filter
    st.session_state["_p_ef_sort_col"] = sort_col
    st.session_state["_p_ef_sort_dir"] = sort_dir

    if clear_filters:
        for k in ["ef_tipo", "ef_cesfam", "ef_planta", "ef_cargo", "ef_unidad",
                   "editor_search", "editor_rev_filter", "ef_sort_col", "ef_sort_dir",
                   "_p_ef_tipo", "_p_ef_cesfam", "_p_ef_planta", "_p_ef_cargo",
                   "_p_ef_unidad", "_p_ef_search", "_p_ef_rev_filter",
                   "_p_ef_sort_col", "_p_ef_sort_dir"]:
            st.session_state.pop(k, None)
        st.rerun()

    # ── Aplicar filtros ──
    dff = df.copy()

    # Filtros de sidebar (globales)
    dff = apply_main_filters(dff, filtros)

    # Filtros inline del editor (sin estado revisión ni búsqueda)
    if f_tipo:
        dff = dff[dff["TIPO"].isin(f_tipo)]
    if f_cesfam:
        dff = dff[dff["CESFAM"].isin(f_cesfam)]
    if f_planta:
        dff = dff[dff["DESCRIPCIÓN DE PLANTA"].isin(f_planta)]
    if f_cargo:
        dff = dff[dff["CARGO"].isin(f_cargo)]
    if f_ud:
        dff = dff[dff["UNIDAD DE DESEMPEÑO"].isin(f_ud)]

    # Snapshot ANTES de filtrar por revisión/búsqueda → base para KPIs de avance
    dff_base = dff.copy()

    # Estado revisión
    if rev_filter == "✅ Revisados":
        dff = dff[dff["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"])]
    elif rev_filter == "⏳ Pendientes":
        dff = dff[~dff["REVISADO"].str.strip().str.upper().isin(["OK", "SI", "SÍ", "X", "✓"])]

    # Búsqueda textual
    if search_text and search_text.strip():
        q = search_text.strip().upper()
        mask = pd.Series(False, index=dff.index)
        for col in ["RUT", "NOMBRE PROFESIONAL", "CESFAM", "CARGO", "ENCOMENDACIONES",
                     "OBSERVACIONES", "UNIDAD DE DESEMPEÑO"]:
            if col in dff.columns:
                mask = mask | dff[col].fillna("").str.upper().str.contains(q, regex=False)
        dff = dff[mask]

    # Ordenamiento
    if sort_col != "(sin ordenar)" and sort_col in dff.columns:
        ascending = sort_dir.startswith("↑")
        dff = dff.sort_values(sort_col, ascending=ascending, na_position="last")

    # ── KPIs de avance (sobre dff_base = filtro completo sin revisión/búsqueda) ──
    kpis_base = compute_kpis(dff_base)
    total_base = kpis_base["total"]
    rev_base = kpis_base["revisados"]
    pend_base = kpis_base["pendientes"]
    pct_base = kpis_base["pct_revisados"]
    n_filters = sum(1 for x in [f_tipo, f_cesfam, f_planta, f_cargo, f_ud] if x) + (1 if search_text else 0) + (1 if rev_filter != "Todos" else 0)

    # Contexto: nombres de CESFAM filtrados
    _cesfam_label = ", ".join(f_cesfam) if f_cesfam else "Todos los CESFAM"

    # Barra de progreso visual
    _bar_pct = min(pct_base, 100)
    if _bar_pct >= 80:
        _bar_color = "#22C55E"
    elif _bar_pct >= 40:
        _bar_color = "#FACC15"
    else:
        _bar_color = "#EF4444"

    st.markdown(f"""<div style="background:#F8FAFC; border:2px solid #E2E8F0; border-radius:10px;
        padding:.8rem 1.2rem; margin-bottom:.6rem;">
        <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:.4rem;">
            <span style="font-weight:700; color:#0A2E1F; font-size:.95rem;">📍 {_cesfam_label}</span>
            <span style="font-size:.82rem; color:#64748B;">{n_filters} filtro(s) activo(s) · Mostrando {len(dff):,} de {total_base:,} registros</span>
        </div>
        <div style="display:flex; gap:1.5rem; align-items:center; margin-bottom:.5rem;">
            <div style="text-align:center; min-width:90px;">
                <div style="font-size:1.6rem; font-weight:800; color:#0A2E1F;">{total_base:,}</div>
                <div style="font-size:.72rem; color:#64748B; text-transform:uppercase; letter-spacing:.5px;">Total</div>
            </div>
            <div style="text-align:center; min-width:90px;">
                <div style="font-size:1.6rem; font-weight:800; color:#22C55E;">{rev_base:,}</div>
                <div style="font-size:.72rem; color:#64748B; text-transform:uppercase; letter-spacing:.5px;">Revisados</div>
            </div>
            <div style="text-align:center; min-width:90px;">
                <div style="font-size:1.6rem; font-weight:800; color:#EF4444;">{pend_base:,}</div>
                <div style="font-size:.72rem; color:#64748B; text-transform:uppercase; letter-spacing:.5px;">Pendientes</div>
            </div>
            <div style="flex:1; padding:0 .5rem;">
                <div style="display:flex; justify-content:space-between; margin-bottom:2px;">
                    <span style="font-size:.78rem; font-weight:600; color:#0A2E1F;">Avance</span>
                    <span style="font-size:.78rem; font-weight:700; color:{_bar_color};">{pct_base:.1f}%</span>
                </div>
                <div style="background:#E2E8F0; border-radius:6px; height:14px; overflow:hidden;">
                    <div style="background:{_bar_color}; width:{_bar_pct:.1f}%; height:100%; border-radius:6px;
                        transition: width 0.3s ease;"></div>
                </div>
                <div style="font-size:.7rem; color:#94A3B8; margin-top:2px; text-align:right;">
                    {"✅ ¡Completado!" if pct_base >= 100 else f"Faltan {pend_base:,} registros por revisar"}
                </div>
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    st.markdown('<div class="section-title">Tabla Editable — Doble clic en celda para editar · Seleccione fila para asignar encomendaciones</div>', unsafe_allow_html=True)

    # Columnas a mostrar
    display_cols = ["REVISADO", "TIPO", "RUT", "CESFAM", "NOMBRE PROFESIONAL",
                    "DESCRIPCIÓN DE PLANTA", "CARGO", "UNIDAD DE DESEMPEÑO",
                    "ENCOMENDACIONES", "OBSERVACIONES", "Horas por contrato",
                    "Horas Totales", "Total Descuentos semanal (horas)", "Total Horas Clínicas"]
    all_cols = [c for c in display_cols if c in dff.columns]

    # Contar personas nuevas
    n_nuevos = int(dff["_ES_NUEVO"].sum()) if "_ES_NUEVO" in dff.columns else 0
    if n_nuevos > 0:
        st.markdown(f'<div class="warn-banner">🆕 {n_nuevos} persona(s) nueva(s) importada(s) desde Excel — resaltadas en la tabla</div>', unsafe_allow_html=True)


    # ── Paginación real ──
    PAGE_SIZE = 100
    total_rows = len(dff)
    total_pages = (total_rows - 1) // PAGE_SIZE + 1
    page_num = st.session_state.get("_editor_page", 1)
    if page_num < 1:
        page_num = 1
    if page_num > total_pages:
        page_num = total_pages
    start_idx = (page_num - 1) * PAGE_SIZE
    end_idx = start_idx + PAGE_SIZE
    dff_page = dff.iloc[start_idx:end_idx]

    # Navegación de páginas
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("⬅️ Anterior", disabled=(page_num == 1)):
            st.session_state["_editor_page"] = page_num - 1
            st.rerun()
    with col2:
        st.markdown(f"<div style='text-align:center; font-size:0.95rem; margin-top:0.5rem;'>Página <b>{page_num}</b> de <b>{total_pages}</b> ({total_rows:,} registros)</div>", unsafe_allow_html=True)
    with col3:
        if st.button("Siguiente ➡️", disabled=(page_num == total_pages)):
            st.session_state["_editor_page"] = page_num + 1
            st.rerun()

    dff_edit = dff_page[all_cols].copy()
    dff_edit["_real_idx"] = dff_page.index
    if "_ES_NUEVO" in dff.columns:
        dff_edit.insert(0, "🆕", dff_page["_ES_NUEVO"].map({True: "🆕", False: ""}).fillna(""))
    else:
        dff_edit.insert(0, "🆕", "")

    # ── Configurar AG Grid ──
    gb = GridOptionsBuilder.from_dataframe(dff_edit)
    gb.configure_default_column(editable=True, resizable=True, sortable=True, filter=True, minWidth=80)

    # Columnas fijas a la izquierda (frozen)
    gb.configure_column("🆕", pinned="left", width=60, editable=False, headerName="Nuevo")
    gb.configure_column("REVISADO", pinned="left", width=80, headerName="Rev",
                        cellEditor="agSelectCellEditor",
                        cellEditorParams={"values": ["", "OK"]})
    gb.configure_column("NOMBRE PROFESIONAL", pinned="left", width=220, headerName="Nombre")
    gb.configure_column("TIPO", pinned="left", width=70, headerName="Tipo")
    gb.configure_column("RUT", pinned="left", width=100, headerName="RUT")

    # Columnas no fijas
    gb.configure_column("CESFAM", width=140, headerName="CESFAM")
    gb.configure_column("CARGO", width=160, headerName="Cargo")
    gb.configure_column("DESCRIPCIÓN DE PLANTA", width=140, headerName="Planta")
    gb.configure_column("UNIDAD DE DESEMPEÑO", width=160, headerName="Unidad")
    gb.configure_column("ENCOMENDACIONES", width=300, headerName="Encomendaciones")
    gb.configure_column("OBSERVACIONES", width=200, headerName="Obs.")
    _fmt1 = JsCode("function(p){return p.value!=null? Number(p.value).toFixed(1): ''}")
    gb.configure_column("Horas por contrato", width=100, headerName="Hrs Contr.",
                        type=["numericColumn"], valueParser="Number(newValue)",
                        valueFormatter=_fmt1)
    gb.configure_column("Horas Totales", width=100, headerName="Hrs Tot.",
                        type=["numericColumn"], valueParser="Number(newValue)",
                        valueFormatter=_fmt1)
    gb.configure_column("Total Descuentos semanal (horas)", width=100, headerName="Desc.",
                        type=["numericColumn"], valueParser="Number(newValue)",
                        valueFormatter=_fmt1)
    gb.configure_column("Total Horas Clínicas", width=100, headerName="Hrs Clín.",
                        type=["numericColumn"], valueParser="Number(newValue)",
                        valueFormatter=_fmt1)
    # Ocultar columna auxiliar de índice
    gb.configure_column("_real_idx", hide=True)

    # Selección de fila con checkbox para encomendaciones
    gb.configure_selection("single", use_checkbox=True)

    grid_options = gb.build()

    # CSS corporativo para los encabezados AG Grid
    custom_css = {
        ".ag-header-cell": {
            "background": "linear-gradient(135deg, #0A2E1F, #145A38) !important",
        },
        ".ag-header-cell-label": {
            "color": "#4ADE80 !important",
            "font-weight": "700 !important",
            "font-size": "0.78rem !important",
            "letter-spacing": "0.3px",
        },
        ".ag-header": {
            "border-bottom": "2px solid #4ADE80 !important",
        },
        ".ag-row-selected": {
            "background-color": "#F0FDF4 !important",
            "border-left": "3px solid #4ADE80 !important",
        },
        ".ag-root-wrapper": {
            "border": "2px solid #0A2E1F !important",
            "border-radius": "8px !important",
        },
    }


    response = AgGrid(
        dff_edit,
        gridOptions=grid_options,
        update_mode=GridUpdateMode.MODEL_CHANGED,
        data_return_mode=DataReturnMode.AS_INPUT,
        height=500,
        theme="alpine",
        custom_css=custom_css,
        allow_unsafe_jscode=True,
        key=f"ag_main_{st.session_state.get('_ag_version', 0)}",
    )

    # ── 1. Procesar selección de fila (actualiza _sel_rut AHORA) ──
    selected = response.selected_rows if hasattr(response, "selected_rows") else None
    if selected is not None:
        if isinstance(selected, pd.DataFrame) and len(selected) > 0:
            sel_rut = str(selected.iloc[0].get("RUT", "")).strip()
            if sel_rut:
                st.session_state["_sel_rut"] = sel_rut
        elif isinstance(selected, list) and len(selected) > 0:
            sel_rut = str(selected[0].get("RUT", "")).strip()
            if sel_rut:
                st.session_state["_sel_rut"] = sel_rut

    # ── 2. Sincronizar ediciones de celdas (comparar solo contra dff_page) ──
    edited_df = response.data if hasattr(response, "data") else None
    if isinstance(edited_df, pd.DataFrame) and not edited_df.empty:
        n_compare = min(len(dff_page), len(edited_df))
        for col in all_cols:
            if col not in edited_df.columns:
                continue
            for i in range(n_compare):
                real_idx_i = dff_page.index[i]
                sent_val = dff_page.iloc[i][col]
                grid_val = edited_df.iloc[i][col]
                sent_s = str(sent_val) if pd.notna(sent_val) else ""
                grid_s = str(grid_val) if pd.notna(grid_val) else ""
                if grid_s != sent_s:
                    st.session_state.df_main.at[real_idx_i, col] = grid_val
                    if col == "ENCOMENDACIONES":
                        lookup = st.session_state.horas_lookup
                        items = [x.strip() for x in grid_s.split("|") if x.strip()]
                        td = sum(lookup.get(it, 0) for it in items)
                        ht = float(st.session_state.df_main.at[real_idx_i, "Horas Totales"]) if pd.notna(st.session_state.df_main.at[real_idx_i, "Horas Totales"]) else 0
                        st.session_state.df_main.at[real_idx_i, "Total Descuentos semanal (horas)"] = td
                        st.session_state.df_main.at[real_idx_i, "Total Horas Clínicas"] = ht - td
                    st.session_state.dirty_main = True

    # Guardar cambios inline sin reiniciar el grid (no incrementar _ag_version)
    if st.session_state.dirty_main:
        db_save_main(st.session_state.df_main)
        st.session_state.dirty_main = False

    # ── 3. Editor de encomendaciones — DESPUÉS del AgGrid con _sel_rut ya actualizado ──
    st.markdown('<div class="section-title">Encomendaciones — Fila seleccionada</div>', unsafe_allow_html=True)
    _active_rut = st.session_state.get("_sel_rut", None)
    _row = None
    _real_idx = None
    if _active_rut:
        _match = df["RUT"].astype(str).str.strip() == _active_rut
        if _match.any():
            _real_idx = df[_match].index[0]
            _row = df.loc[_real_idx]
        else:
            st.session_state.pop("_sel_rut", None)

    if _row is not None:
        _nombre = str(_row.get("NOMBRE PROFESIONAL", "")) if pd.notna(_row.get("NOMBRE PROFESIONAL")) else ""
        _rut    = str(_row.get("RUT", "")) if pd.notna(_row.get("RUT")) else ""
        _cargo  = str(_row.get("CARGO", "")) if pd.notna(_row.get("CARGO")) else ""
        _cesfam = str(_row.get("CESFAM", "")) if pd.notna(_row.get("CESFAM")) else ""
        _hrs_tot = float(_row.get("Horas Totales", 0)) if pd.notna(_row.get("Horas Totales")) else 0
        _lookup = st.session_state.horas_lookup
        _activity_names = sorted(_lookup.keys())
        _cur_encom = str(_row.get("ENCOMENDACIONES", "")) if pd.notna(_row.get("ENCOMENDACIONES")) else ""
        _cur_items = [x.strip() for x in _cur_encom.split("|") if x.strip()] if _cur_encom.strip() else []

        st.markdown(f"""<div style="background:#F0FDF4; border:1px solid #4ADE80; padding:.5rem .8rem;
                    border-radius:6px; margin:.3rem 0 .5rem 0;">
            <span style="color:#0A2E1F; font-weight:600;">✏️ {_nombre}</span>
            <span style="color:#666; font-size:.82rem;"> · {_rut} · {_cargo} · {_cesfam} · {_hrs_tot:.0f} hrs</span>
        </div>""", unsafe_allow_html=True)

        _new_items = st.multiselect(
            "🔍 Encomendaciones (escriba para buscar actividades)",
            options=_activity_names,
            default=[x for x in _cur_items if x in _activity_names],
            key=f"encom_ms_{_real_idx}",
            placeholder="Escriba para buscar y agregar actividades...",
        )
        _total_desc = sum(_lookup.get(item, 0) for item in _new_items)
        _hrs_cli = _hrs_tot - _total_desc
        _jornadas = _hrs_cli / 44 if _hrs_cli > 0 else 0
        _info_col, _btn_col = st.columns([4, 1])
        with _info_col:
            st.caption(f"Descuentos: {_total_desc:.2f} · Hrs Clínicas: {_hrs_cli:.2f} · Jornadas: {_jornadas:.2f} · Actividades: {len(_new_items)}")
        with _btn_col:
            _new_encom_str = " | ".join(sorted(_new_items))
            _old_sorted    = " | ".join(sorted(_cur_items))
            if _new_encom_str != _old_sorted:
                if st.button("✅ Aplicar", type="primary", key=f"apply_{_real_idx}"):
                    st.session_state.df_main.at[_real_idx, "ENCOMENDACIONES"] = _new_encom_str
                    st.session_state.df_main.at[_real_idx, "Total Descuentos semanal (horas)"] = _total_desc
                    st.session_state.df_main.at[_real_idx, "Total Horas Clínicas"] = _hrs_cli
                    st.session_state.df_main.at[_real_idx, "REVISADO"] = "OK"
                    db_save_main(st.session_state.df_main)
                    st.session_state["_ag_version"] = st.session_state.get("_ag_version", 0) + 1
                    st.toast(f"Encomendaciones de {_nombre} guardadas", icon="💾")
                    st.rerun()
    else:
        st.markdown("""<div style="background:#F0FDF4; border:1px solid #BBF7D0; padding:.6rem 1rem;
            border-radius:8px; color:#666; font-size:.85rem;">
            ℹ️ Seleccione una fila en la tabla para editar encomendaciones.
        </div>""", unsafe_allow_html=True)


def recalculate_hours(df: pd.DataFrame):
    """Recalcula Total Descuentos y Total Horas Clínicas basado en ENCOMENDACIONES."""
    lookup = st.session_state.horas_lookup

    def _calc_desc(encom_str):
        if not encom_str or not str(encom_str).strip():
            return 0.0
        return sum(lookup.get(item.strip(), 0) for item in str(encom_str).split("|") if item.strip())

    df["Total Descuentos semanal (horas)"] = df["ENCOMENDACIONES"].fillna("").apply(_calc_desc)
    hrs_tot = pd.to_numeric(df["Horas Totales"], errors="coerce").fillna(0)
    df["Total Horas Clínicas"] = hrs_tot - df["Total Descuentos semanal (horas)"]


# ─────────────────────────────────────────────────────────────
# PÁGINA: HORAS INDIRECTAS
# ─────────────────────────────────────────────────────────────
def page_horas():
    st.markdown("""
    <div class="main-header">
        <h1>⏱️ Tabla de Horas Indirectas</h1>
        <p>Edite los valores de horas por actividad · Se actualizan en tiempo real los descuentos</p>
    </div>
    """, unsafe_allow_html=True)

    df_horas = st.session_state.df_horas
    if df_horas is None:
        st.warning("No se han cargado las Horas Indirectas. Suba el archivo Excel primero.")
        return

    st.markdown('<div class="section-title">Tabla Editable — Horas Indirectas</div>', unsafe_allow_html=True)
    st.caption("Modifique 'Horas a la semana' para actualizar los cálculos automáticamente.")

    # Determinar columnas hasta "Horas a la semana"
    _hrs_sem_col = None
    for _c in df_horas.columns:
        if "semana" in str(_c).lower():
            _hrs_sem_col = _c
            break
    if _hrs_sem_col:
        _col_idx = list(df_horas.columns).index(_hrs_sem_col)
        _display_cols_h = list(df_horas.columns[:_col_idx + 1])
    else:
        _display_cols_h = list(df_horas.columns)
    df_horas_display = df_horas[_display_cols_h]

    column_config = {
        df_horas_display.columns[0]: st.column_config.TextColumn("Actividad", width="large"),
    }
    for c in df_horas_display.columns[1:]:
        if c == _hrs_sem_col:
            column_config[c] = st.column_config.NumberColumn(c, format="%.1f")
        else:
            column_config[c] = st.column_config.NumberColumn(c, format="%.0f")

    edited = st.data_editor(
        df_horas_display,
        column_config=column_config,
        disabled=[df_horas_display.columns[0]],
        use_container_width=True,
        height=700,
        num_rows="fixed",
        key="editor_horas",
    )

    if edited is not None:
        # Detectar cambios
        changed = False
        for col in df_horas_display.columns[1:]:
            for i in range(len(df_horas_display)):
                o = df_horas.at[i, col]
                n = edited.at[i, col]
                o_f = float(o) if pd.notna(o) else 0
                n_f = float(n) if pd.notna(n) else 0
                if abs(o_f - n_f) > 0.0001:
                    st.session_state.df_horas.at[i, col] = n
                    changed = True
        if changed:
            st.session_state.horas_lookup = build_horas_lookup(st.session_state.df_horas)
            recalculate_hours(st.session_state.df_main)
            # Auto-guardar en SQLite
            db_save_main(st.session_state.df_main)
            db_save_horas(st.session_state.df_horas)
            try:
                save_horas_sheet(st.session_state.excel_path, st.session_state.df_horas)
            except Exception:
                pass
            st.session_state.dirty_horas = False
            st.session_state.dirty_main = False

    # Resumen visual
    st.markdown('<div class="section-title">Resumen de Horas por Categoría</div>', unsafe_allow_html=True)
    lookup = st.session_state.horas_lookup
    if lookup:
        items = sorted(lookup.items(), key=lambda x: x[1], reverse=True)[:20]
        df_viz = pd.DataFrame(items, columns=["Actividad", "Horas/semana"])
        fig = px.bar(df_viz, x="Horas/semana", y="Actividad", orientation="h",
                     color_discrete_sequence=[C_ACCENT], text="Horas/semana")
        fig.update_layout(height=max(300, len(df_viz) * 28), margin=dict(l=0,r=20,t=10,b=0),
                          font=dict(family="Inter", size=11))
        fig.update_traces(textposition="outside", texttemplate="%{text:.1f}")
        st.plotly_chart(fig, use_container_width=True)


# ─────────────────────────────────────────────────────────────
# PÁGINA: DOT IDEAL
# ─────────────────────────────────────────────────────────────
def page_dot_ideal():
    st.markdown("""
    <div class="main-header">
        <h1>🎯 Dotación Ideal por Centro</h1>
        <p>Tabla editable de jornadas ideales por CESFAM y cargo profesional</p>
    </div>
    """, unsafe_allow_html=True)

    df_dot = st.session_state.df_dot
    if df_dot is None or df_dot.empty:
        st.warning("No se pudo cargar la tabla DOT IDEAL del Excel.")
        return

    st.markdown('<div class="section-title">Dotación Ideal — Jornadas por CESFAM</div>', unsafe_allow_html=True)

    # Columnas de jornadas profesionales (decimales)
    _jornada_cols = {"MEDICOS", "ODONT.", "TONS", "QF", "EU", "NUTRI", "MATRONA", "PS", "TS", "TENS", "ADM"}
    _int_cols_dot = [c for c in df_dot.columns if c not in ["CESFAM", "CECOSF"] and c not in _jornada_cols]

    # Crear copia de display con separador de miles en columnas enteras
    df_dot_display = df_dot.copy()
    for _c in _int_cols_dot:
        if _c in df_dot_display.columns:
            df_dot_display[_c] = df_dot_display[_c].apply(
                lambda x: f"{int(x):,}".replace(",", ".") if pd.notna(x) and isinstance(x, (int, float)) else (str(x) if pd.notna(x) else "")
            )

    column_config = {
        "CESFAM": st.column_config.TextColumn("CESFAM", width="large"),
    }
    for c in _int_cols_dot:
        column_config[c] = st.column_config.TextColumn(c)
    for c in _jornada_cols:
        if c in df_dot_display.columns:
            column_config[c] = st.column_config.NumberColumn(c, format="%.2f")

    _disabled_cols = ["CESFAM", "CECOSF"] + _int_cols_dot

    edited_dot = st.data_editor(
        df_dot_display,
        column_config=column_config,
        disabled=_disabled_cols,
        use_container_width=True,
        height=500,
        num_rows="fixed",
        key="editor_dot",
    )

    if edited_dot is not None:
        # Solo sincronizar columnas de jornadas editables al df original
        for _jc in _jornada_cols:
            if _jc in edited_dot.columns and _jc in st.session_state.df_dot.columns:
                st.session_state.df_dot[_jc] = edited_dot[_jc]
        db_save_dot(st.session_state.df_dot)

    # Heatmap visual
    st.markdown('<div class="section-title">Mapa de Calor — Jornadas Ideales</div>', unsafe_allow_html=True)
    display_dot = st.session_state.df_dot.copy()
    if "CESFAM" in display_dot.columns:
        display_dot = display_dot.set_index("CESFAM")
        # Solo columnas numéricas de cargos
        cargo_cols = ["MEDICOS", "ODONT.", "TONS", "QF", "EU", "NUTRI", "MATRONA", "PS", "TS", "TENS", "ADM"]
        available = [c for c in cargo_cols if c in display_dot.columns]
        if available:
            heat_df = display_dot[available].apply(pd.to_numeric, errors="coerce").fillna(0)
            fig = px.imshow(heat_df, text_auto=".1f", aspect="auto",
                            color_continuous_scale=[[0, "#F0FDF4"], [0.5, "#86EFAC"], [1, "#0A2E1F"]])
            fig.update_layout(height=400, margin=dict(l=0,r=0,t=10,b=0),
                              font=dict(family="Inter", size=11))
            st.plotly_chart(fig, use_container_width=True)


# ─────────────────────────────────────────────────────────────
# PÁGINA: ANÁLISIS DE BRECHAS
# ─────────────────────────────────────────────────────────────
def page_brechas(filtros):
    st.markdown("""
    <div class="main-header">
        <h1>📈 Panorama Global de Brechas — Jornadas de Atención Directa</h1>
        <p>Distribución de jornadas reales vs ideales por CESFAM · Herramienta para redistribución de horas</p>
    </div>
    """, unsafe_allow_html=True)

    df = st.session_state.df_main
    df_dot = st.session_state.df_dot

    if df_dot is None or df_dot.empty:
        st.warning("No se pudo cargar la tabla DOT IDEAL.")
        return

    if "CESFAM" not in df.columns or "CARGO" not in df.columns:
        st.error("Faltan columnas CESFAM o CARGO en los datos.")
        return

    dot_categories = ["MEDICOS", "ODONT.", "EU", "NUTRI", "MATRONA", "PS", "TS", "QF", "TENS", "ADM"]

    # ═══════════════════════════════════════════════════════════
    # CONSTRUIR MATRIZ: CESFAM × Estamento → jornadas reales, ideales, brecha
    # ═══════════════════════════════════════════════════════════
    cesfam_list = sorted(df["CESFAM"].dropna().unique().tolist())
    rows_global = []         # una fila por CESFAM (totales)
    rows_detail = []         # una fila por CESFAM × estamento

    for cesfam_val in cesfam_list:
        dfc = df[df["CESFAM"] == cesfam_val]
        matched_cesfam = match_cesfam_to_dot(cesfam_val, df_dot)
        dot_row = None
        if matched_cesfam:
            dot_match = df_dot[df_dot["CESFAM"].astype(str).str.strip() == matched_cesfam]
            if not dot_match.empty:
                dot_row = dot_match.iloc[0]

        total_jr_real = 0
        total_jr_ideal = 0
        has_ideal = False

        for dot_col in dot_categories:
            cargos_map = [k for k, v in CARGO_DOT_MAP.items() if v == dot_col]
            mask = dfc["CARGO"].str.strip().isin(cargos_map)
            if not mask.any():
                mask = dfc["CARGO"].str.strip().str.upper().isin([c.upper() for c in cargos_map])
            sub = dfc[mask]

            hrs_cli = sub["Total Horas Clínicas"].sum() if not sub.empty and "Total Horas Clínicas" in sub.columns else 0
            jr_real = hrs_cli / 44 if hrs_cli > 0 else 0

            jr_ideal = None
            if dot_row is not None and dot_col in dot_row.index:
                v = dot_row[dot_col]
                if pd.notna(v):
                    jr_ideal = float(v)

            brecha = None
            if jr_ideal is not None:
                brecha = jr_real - jr_ideal
                has_ideal = True
                total_jr_ideal += jr_ideal

            total_jr_real += jr_real

            rows_detail.append({
                "CESFAM": cesfam_val,
                "Estamento": dot_col,
                "Jornadas Reales": round(jr_real, 2),
                "Jornadas Ideales": round(jr_ideal, 2) if jr_ideal is not None else None,
                "Brecha": round(brecha, 2) if brecha is not None else None,
            })

        # Funcionarios no mapeados
        all_mapped_cargos = set(CARGO_DOT_MAP.keys()) | {k.upper() for k in CARGO_DOT_MAP.keys()}
        unmapped = dfc[~dfc["CARGO"].str.strip().isin(all_mapped_cargos) &
                       ~dfc["CARGO"].str.strip().str.upper().isin(all_mapped_cargos)]
        if not unmapped.empty:
            hrs_cli_otros = unmapped["Total Horas Clínicas"].sum() if "Total Horas Clínicas" in unmapped.columns else 0
            jr_otros = hrs_cli_otros / 44 if hrs_cli_otros > 0 else 0
            total_jr_real += jr_otros
            rows_detail.append({
                "CESFAM": cesfam_val,
                "Estamento": "OTROS",
                "Jornadas Reales": round(jr_otros, 2),
                "Jornadas Ideales": None,
                "Brecha": None,
            })

        brecha_total = (total_jr_real - total_jr_ideal) if has_ideal else None
        n_func = len(dfc)
        poblacion = ""
        if dot_row is not None and "POBLACION" in dot_row.index and pd.notna(dot_row["POBLACION"]):
            poblacion = int(dot_row["POBLACION"])

        rows_global.append({
            "CESFAM": cesfam_val,
            "Funcionarios": n_func,
            "Población": poblacion if poblacion else "—",
            "Jornadas Reales": round(total_jr_real, 2),
            "Jornadas Ideales": round(total_jr_ideal, 2) if has_ideal else "—",
            "Brecha (Jornadas)": round(brecha_total, 2) if brecha_total is not None else "—",
            "_brecha_num": brecha_total if brecha_total is not None else 0,
            "_jr_real": total_jr_real,
            "_jr_ideal": total_jr_ideal if has_ideal else 0,
        })

    df_global = pd.DataFrame(rows_global)
    df_detail = pd.DataFrame(rows_detail)

    # ═══════════════════════════════════════════════════════════
    # KPIs GLOBALES
    # ═══════════════════════════════════════════════════════════
    sum_jr_real = df_global["_jr_real"].sum()
    sum_jr_ideal = df_global["_jr_ideal"].sum()
    sum_brecha = df_global["_brecha_num"].sum()
    n_deficit = sum(1 for r in rows_global if isinstance(r["Brecha (Jornadas)"], (int, float)) and r["Brecha (Jornadas)"] < -0.5)
    n_superavit = sum(1 for r in rows_global if isinstance(r["Brecha (Jornadas)"], (int, float)) and r["Brecha (Jornadas)"] > 0.5)
    n_equilibrio = sum(1 for r in rows_global if isinstance(r["Brecha (Jornadas)"], (int, float)) and abs(r["Brecha (Jornadas)"]) <= 0.5)

    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#0A2E1F,#145A38);color:#fff;padding:1rem 1.5rem;
        border-radius:10px;margin-bottom:1rem;">
        <div style="display:flex;gap:2rem;flex-wrap:wrap;align-items:center;">
            <div style="text-align:center;min-width:130px;">
                <div style="font-size:2rem;font-weight:800;">{sum_jr_real:.1f}</div>
                <div style="font-size:.72rem;text-transform:uppercase;letter-spacing:.5px;color:#A8D8B9;">
                    Jornadas Reales Totales</div>
            </div>
            <div style="text-align:center;min-width:130px;">
                <div style="font-size:2rem;font-weight:800;">{sum_jr_ideal:.1f}</div>
                <div style="font-size:.72rem;text-transform:uppercase;letter-spacing:.5px;color:#A8D8B9;">
                    Jornadas Ideales Totales</div>
            </div>
            <div style="text-align:center;min-width:130px;">
                <div style="font-size:2rem;font-weight:800;color:{'#4ADE80' if sum_brecha >= 0 else '#FF6B6B'};">
                    {sum_brecha:+.1f}</div>
                <div style="font-size:.72rem;text-transform:uppercase;letter-spacing:.5px;color:#A8D8B9;">
                    Brecha Neta (jornadas)</div>
            </div>
            <div style="border-left:1px solid #4ADE80;padding-left:1.5rem;display:flex;gap:1.5rem;">
                <div style="text-align:center;">
                    <div style="font-size:1.5rem;font-weight:700;color:#FF6B6B;">{n_deficit}</div>
                    <div style="font-size:.7rem;color:#ccc;">Con Déficit</div>
                </div>
                <div style="text-align:center;">
                    <div style="font-size:1.5rem;font-weight:700;color:#4ADE80;">{n_equilibrio}</div>
                    <div style="font-size:.7rem;color:#ccc;">Equilibrio</div>
                </div>
                <div style="text-align:center;">
                    <div style="font-size:1.5rem;font-weight:700;color:#FACC15;">{n_superavit}</div>
                    <div style="font-size:.7rem;color:#ccc;">Con Superávit</div>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════
    # 1. GRÁFICO WATERFALL — Brecha por CESFAM
    # ═══════════════════════════════════════════════════════════
    st.markdown('<div class="section-title">Brecha Neta por CESFAM (Jornadas de Atención Directa)</div>',
                unsafe_allow_html=True)
    st.caption("Barras hacia la derecha = superávit (sobran jornadas) · Barras hacia la izquierda = déficit (faltan jornadas)")

    df_chart = df_global[df_global["_brecha_num"] != 0].sort_values("_brecha_num")
    if not df_chart.empty:
        colors = [C_DANGER if v < -0.5 else (C_SUCCESS if v > 0.5 else C_ACCENT)
                  for v in df_chart["_brecha_num"]]

        fig_wf = go.Figure()
        fig_wf.add_trace(go.Bar(
            y=df_chart["CESFAM"],
            x=df_chart["_brecha_num"],
            orientation="h",
            marker_color=colors,
            text=[f"{v:+.1f}" for v in df_chart["_brecha_num"]],
            textposition="outside",
            textfont=dict(size=11, color="#333"),
        ))
        fig_wf.add_vline(x=0, line_dash="solid", line_color="#666", line_width=1.5)
        fig_wf.update_layout(
            height=max(350, len(df_chart) * 42),
            margin=dict(l=10, r=60, t=10, b=10),
            xaxis_title="Brecha (jornadas): ← déficit | superávit →",
            yaxis=dict(autorange="reversed"),
            plot_bgcolor="#FAFAFA",
            font=dict(family="Inter, Calibri", size=11),
        )
        st.plotly_chart(fig_wf, use_container_width=True)
    else:
        st.info("No hay brechas calculadas (verifique la tabla DOT IDEAL).")

    # ═══════════════════════════════════════════════════════════
    # 2. TABLA RESUMEN GLOBAL — CESFAM con Real vs Ideal
    # ═══════════════════════════════════════════════════════════
    st.markdown('<div class="section-title">Tabla Resumen por CESFAM</div>', unsafe_allow_html=True)

    df_display = df_global.drop(columns=["_brecha_num", "_jr_real", "_jr_ideal"])

    def _style_brecha_row(row):
        styles = [""] * len(row)
        brecha_val = row.get("Brecha (Jornadas)")
        if isinstance(brecha_val, (int, float)):
            if brecha_val < -0.5:
                styles = ["background-color: #FEE2E2"] * len(row)  # rojo suave
            elif brecha_val > 0.5:
                styles = ["background-color: #DCFCE7"] * len(row)  # verde suave
            else:
                styles = ["background-color: #F0FDF4"] * len(row)
        return styles

    _fmt_miles = lambda x: f"{int(x):,}".replace(",", ".") if isinstance(x, (int, float)) and pd.notna(x) else str(x)
    _fmt_dec1 = lambda x: f"{x:.1f}" if isinstance(x, (int, float)) and pd.notna(x) else str(x)

    st.dataframe(
        df_display.style
            .apply(_style_brecha_row, axis=1)
            .format({
                "Funcionarios": _fmt_miles,
                "Población": _fmt_miles,
                "Jornadas Reales": _fmt_dec1,
                "Jornadas Ideales": _fmt_dec1,
                "Brecha (Jornadas)": _fmt_dec1,
            }),
        use_container_width=True,
        height=min(600, 40 + 35 * len(df_display)),
    )

    # ═══════════════════════════════════════════════════════════
    # 3. GRÁFICO STACKED — Real vs Ideal por CESFAM
    # ═══════════════════════════════════════════════════════════
    st.markdown('<div class="section-title">Jornadas Reales vs Ideales por CESFAM</div>',
                unsafe_allow_html=True)
    st.caption("Comparación directa: barra verde = jornadas reales, línea = jornadas ideales")

    df_bar = df_global.sort_values("_jr_real", ascending=False)
    fig_comp = go.Figure()
    fig_comp.add_trace(go.Bar(
        x=df_bar["CESFAM"],
        y=df_bar["_jr_real"],
        name="Jornadas Reales",
        marker_color="#22C55E",
        text=[f"{v:.1f}" for v in df_bar["_jr_real"]],
        textposition="outside",
        textfont=dict(size=9),
    ))
    fig_comp.add_trace(go.Scatter(
        x=df_bar["CESFAM"],
        y=df_bar["_jr_ideal"],
        name="Jornadas Ideales",
        mode="markers+lines",
        marker=dict(color="#EF4444", size=10, symbol="diamond"),
        line=dict(color="#EF4444", width=2, dash="dash"),
    ))
    fig_comp.update_layout(
        height=450,
        margin=dict(l=0, r=0, t=30, b=10),
        xaxis_tickangle=-45,
        yaxis_title="Jornadas (hrs clínicas / 44)",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        barmode="group",
        plot_bgcolor="#FAFAFA",
        font=dict(family="Inter, Calibri", size=11),
    )
    st.plotly_chart(fig_comp, use_container_width=True)

    # ═══════════════════════════════════════════════════════════
    # 4. HEATMAP — Brecha por CESFAM × Estamento
    # ═══════════════════════════════════════════════════════════
    st.markdown('<div class="section-title">Mapa de Calor: Brechas por CESFAM y Estamento</div>',
                unsafe_allow_html=True)
    st.caption("Rojo = déficit de jornadas · Verde = superávit · Blanco = equilibrio o sin dato")

    # Pivot: CESFAM (filas) × Estamento (columnas) → brecha
    df_pivot = df_detail[df_detail["Estamento"].isin(dot_categories)].pivot_table(
        index="CESFAM", columns="Estamento", values="Brecha", aggfunc="sum"
    )
    # Reordenar columnas según dot_categories
    df_pivot = df_pivot.reindex(columns=[c for c in dot_categories if c in df_pivot.columns])

    if not df_pivot.empty:
        # Calcular total brecha por CESFAM para ordenar
        df_pivot["_TOTAL"] = df_pivot.sum(axis=1, skipna=True)
        df_pivot = df_pivot.sort_values("_TOTAL", ascending=True)
        df_pivot = df_pivot.drop(columns=["_TOTAL"])

        z_vals = df_pivot.values
        z_text = [[f"{v:+.1f}" if pd.notna(v) else "—" for v in row] for row in z_vals]

        fig_hm = go.Figure(data=go.Heatmap(
            z=z_vals,
            x=df_pivot.columns.tolist(),
            y=df_pivot.index.tolist(),
            text=z_text,
            texttemplate="%{text}",
            textfont=dict(size=10),
            colorscale=[
                [0.0, "#EF4444"],    # déficit fuerte → rojo
                [0.4, "#FEE2E2"],    # déficit leve → rosado
                [0.5, "#FFFFFF"],    # equilibrio → blanco
                [0.6, "#DCFCE7"],    # superávit leve → verde claro
                [1.0, "#22C55E"],    # superávit fuerte → verde
            ],
            zmid=0,
            colorbar=dict(title="Brecha<br>(jornadas)", ticksuffix=" jn"),
            hoverongaps=False,
        ))
        fig_hm.update_layout(
            height=max(400, len(df_pivot) * 38),
            margin=dict(l=10, r=10, t=10, b=10),
            xaxis=dict(side="top"),
            yaxis=dict(autorange="reversed"),
            font=dict(family="Inter, Calibri", size=11),
        )
        st.plotly_chart(fig_hm, use_container_width=True)
    else:
        st.info("No hay datos suficientes para el mapa de calor.")

    # ═══════════════════════════════════════════════════════════
    # 5. EXPLORADOR DETALLADO — Seleccionar CESFAM
    # ═══════════════════════════════════════════════════════════
    st.markdown('<div class="section-title">Explorador por CESFAM</div>', unsafe_allow_html=True)

    sel_cesfam = st.selectbox("Seleccione CESFAM para ver detalle por estamento:",
                               cesfam_list, key="brecha_cesfam_sel")
    if sel_cesfam:
        df_det_cesfam = df_detail[df_detail["CESFAM"] == sel_cesfam].copy()
        df_det_cesfam = df_det_cesfam.drop(columns=["CESFAM"])

        # Gráfico de barras agrupadas para ese CESFAM
        df_det_chart = df_det_cesfam[df_det_cesfam["Jornadas Ideales"].notna()].copy()

        if not df_det_chart.empty:
            fig_det = go.Figure()
            fig_det.add_trace(go.Bar(
                x=df_det_chart["Estamento"],
                y=df_det_chart["Jornadas Reales"],
                name="Reales",
                marker_color="#22C55E",
                text=[f"{v:.1f}" for v in df_det_chart["Jornadas Reales"]],
                textposition="outside",
            ))
            fig_det.add_trace(go.Bar(
                x=df_det_chart["Estamento"],
                y=df_det_chart["Jornadas Ideales"],
                name="Ideales",
                marker_color="#3B82F6",
                text=[f"{v:.1f}" for v in df_det_chart["Jornadas Ideales"]],
                textposition="outside",
            ))
            fig_det.update_layout(
                height=400,
                barmode="group",
                margin=dict(l=0, r=0, t=10, b=10),
                yaxis_title="Jornadas",
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                plot_bgcolor="#FAFAFA",
                font=dict(family="Inter, Calibri", size=11),
            )
            st.plotly_chart(fig_det, use_container_width=True)

        # Tabla detalle
        def _style_brecha_detail(row):
            styles = [""] * len(row)
            b = row.get("Brecha")
            if isinstance(b, (int, float)) and pd.notna(b):
                if b < -0.5:
                    styles = ["background-color: #FEE2E2"] * len(row)
                elif b > 0.5:
                    styles = ["background-color: #DCFCE7"] * len(row)
            return styles

        _fmt_d1 = lambda x: f"{x:.1f}" if isinstance(x, (int, float)) and pd.notna(x) else str(x)
        st.dataframe(
            df_det_cesfam.style
                .apply(_style_brecha_detail, axis=1)
                .format({
                    "Jornadas Reales": _fmt_d1,
                    "Jornadas Ideales": _fmt_d1,
                    "Brecha": _fmt_d1,
                }),
            use_container_width=True,
            hide_index=True,
            height=min(500, 40 + 35 * len(df_det_cesfam)),
        )


# ─────────────────────────────────────────────────────────────
# PÁGINA: ACTA DE REUNIÓN
# ─────────────────────────────────────────────────────────────

def _build_acta_data(cesfam: str, df: pd.DataFrame, df_dot) -> dict:
    """Calcula los datos por estamento/cargo para un CESFAM dado."""
    dfc = df[df["CESFAM"] == cesfam].copy()
    total_personas = len(dfc)

    # Agrupar por categoría DOT
    dot_categories = ["MEDICOS", "ODONT.", "EU", "NUTRI", "MATRONA", "PS", "TS", "QF", "TENS", "ADM"]
    estamentos = []
    mapped_indices = set()
    for dot_col in dot_categories:
        # Cargos que mapean a esta categoría
        cargos_map = [k for k, v in CARGO_DOT_MAP.items() if v == dot_col]
        mask = dfc["CARGO"].str.strip().isin(cargos_map)
        if not mask.any():
            mask = dfc["CARGO"].str.strip().str.upper().isin([c.upper() for c in cargos_map])
        sub = dfc[mask]
        if sub.empty:
            continue
        mapped_indices.update(sub.index.tolist())
        n = len(sub)
        hrs_tot = sub["Horas Totales"].sum() if "Horas Totales" in sub.columns else 0
        hrs_desc = sub["Total Descuentos semanal (horas)"].sum() if "Total Descuentos semanal (horas)" in sub.columns else 0
        hrs_cli = sub["Total Horas Clínicas"].sum() if "Total Horas Clínicas" in sub.columns else 0
        jornadas_real = hrs_cli / 44 if hrs_cli > 0 else 0

        # Ideal
        ideal_val = None
        if df_dot is not None and not df_dot.empty:
            matched_cesfam = match_cesfam_to_dot(cesfam, df_dot)
            if matched_cesfam:
                dot_row = df_dot[df_dot["CESFAM"].astype(str).str.strip() == matched_cesfam]
                if not dot_row.empty and dot_col in dot_row.columns:
                    v = dot_row.iloc[0][dot_col]
                    if pd.notna(v):
                        ideal_val = float(v)

        brecha = jornadas_real - ideal_val if ideal_val is not None else None

        estamentos.append({
            "Estamento": dot_col,
            "N° Funcionarios": n,
            "Hrs Totales": round(hrs_tot, 1),
            "Hrs Descuento": round(hrs_desc, 1),
            "Hrs Clínicas": round(hrs_cli, 1),
            "Jornadas Reales": round(jornadas_real, 2),
            "Jornadas Ideales": round(ideal_val, 2) if ideal_val is not None else "—",
            "Brecha": round(brecha, 2) if brecha is not None else "—",
        })

    # Funcionarios con CARGO no mapeado a ninguna categoría DOT
    unmapped = dfc[~dfc.index.isin(mapped_indices)]
    if not unmapped.empty:
        n = len(unmapped)
        hrs_tot = unmapped["Horas Totales"].sum() if "Horas Totales" in unmapped.columns else 0
        hrs_desc = unmapped["Total Descuentos semanal (horas)"].sum() if "Total Descuentos semanal (horas)" in unmapped.columns else 0
        hrs_cli = unmapped["Total Horas Clínicas"].sum() if "Total Horas Clínicas" in unmapped.columns else 0
        jornadas_real = hrs_cli / 44 if hrs_cli > 0 else 0
        estamentos.append({
            "Estamento": "OTROS",
            "N° Funcionarios": n,
            "Hrs Totales": round(hrs_tot, 1),
            "Hrs Descuento": round(hrs_desc, 1),
            "Hrs Clínicas": round(hrs_cli, 1),
            "Jornadas Reales": round(jornadas_real, 2),
            "Jornadas Ideales": "—",
            "Brecha": "—",
        })

    # Detalle por persona
    det_cols = ["RUT", "NOMBRE PROFESIONAL", "CARGO", "Horas Totales",
                "Total Descuentos semanal (horas)", "Total Horas Clínicas", "ENCOMENDACIONES"]
    det_cols = [c for c in det_cols if c in dfc.columns]
    detalle = dfc[det_cols].copy()
    if "Total Horas Clínicas" in detalle.columns:
        detalle["Jornadas"] = (detalle["Total Horas Clínicas"] / 44).round(2)

    # Sectores y CECOSF desde DOT IDEAL
    sectores = ""
    poblacion = ""
    cecosf = ""
    if df_dot is not None and not df_dot.empty:
        matched_cesfam = match_cesfam_to_dot(cesfam, df_dot)
        if matched_cesfam:
            dot_row = df_dot[df_dot["CESFAM"].astype(str).str.strip() == matched_cesfam]
            if not dot_row.empty:
                if "SECTORES" in dot_row.columns:
                    v = dot_row.iloc[0]["SECTORES"]
                    sectores = str(int(v)) if pd.notna(v) else ""
                if "POBLACION" in dot_row.columns:
                    v = dot_row.iloc[0]["POBLACION"]
                    poblacion = f"{int(v):,}" if pd.notna(v) else ""
                if "CECOSF" in dot_row.columns:
                    v = dot_row.iloc[0]["CECOSF"]
                    cecosf = str(int(v)) if pd.notna(v) else ""

    # ── Horas Clínicas por Unidad de Desempeño ──
    unidades_data = []
    if "UNIDAD DE DESEMPEÑO" in dfc.columns:
        for ud, grp in dfc.groupby("UNIDAD DE DESEMPEÑO", dropna=False):
            ud_name = str(ud).strip() if pd.notna(ud) else "(Sin unidad)"
            if not ud_name:
                ud_name = "(Sin unidad)"
            n = len(grp)
            hrs_tot = grp["Horas Totales"].sum() if "Horas Totales" in grp.columns else 0
            hrs_desc = grp["Total Descuentos semanal (horas)"].sum() if "Total Descuentos semanal (horas)" in grp.columns else 0
            hrs_cli = grp["Total Horas Clínicas"].sum() if "Total Horas Clínicas" in grp.columns else 0
            jornadas = hrs_cli / 44 if hrs_cli > 0 else 0
            unidades_data.append({
                "Unidad": ud_name,
                "N° Func.": n,
                "Hrs Totales": round(hrs_tot, 1),
                "Hrs Descuento": round(hrs_desc, 1),
                "Hrs Clínicas": round(hrs_cli, 1),
                "Jornadas": round(jornadas, 2),
            })
        unidades_data.sort(key=lambda x: x["Hrs Clínicas"], reverse=True)

    return {
        "total_personas": total_personas,
        "estamentos": estamentos,
        "unidades": unidades_data,
        "detalle": detalle,
        "sectores": sectores,
        "poblacion": poblacion,
        "cecosf": cecosf,
    }


def _generate_acta_docx(acta_cfg: dict, acta_data: dict) -> bytes:
    """Genera un documento Word institucional con diseño profesional tipo revista."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDocument()

    # ── Colores institucionales SSMC / Dirección APS (tenues para impresión B/N) ──
    _AZUL_GOB    = "4A7BA7"   # Azul gobierno suave
    _ROJO_GOB    = "D4686E"   # Rojo gobierno suave
    _VERDE_SSMC  = "3D7A5F"   # Verde SSMC medio (legible en B/N)
    _VERDE_MEDIO = "5A9E7E"   # Verde medio claro
    _VERDE_CLARO = "A8D8B9"   # Verde acento suave
    _GRIS_OSCURO = "4A5568"
    _GRIS_CLARO  = "F7FAFC"
    _BLANCO      = "FFFFFF"
    _VERDE_FONDO = "EDF7F0"   # Fondo verde muy claro para celdas

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── Estilos base ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(2)

    # ── Helpers ──
    def _shade_cell(cell, color_hex):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell(cell, text, bold=False, size=9, align="left", color=None, font_name="Calibri"):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.name = font_name
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(
                int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
            if val:
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), val.get("val", "single"))
                el.set(qn("w:sz"), val.get("sz", "4"))
                el.set(qn("w:color"), val.get("color", "000000"))
                el.set(qn("w:space"), "0")
                tcBorders.append(el)
        tcPr.append(tcBorders)

    def _add_spacer(height_pt=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(height_pt)

    def _remove_table_borders(table):
        """Remove all borders from a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            borders.append(el)
        tblPr.append(borders)

    # ═══════════════════════════════════════════════════════════
    # ENCABEZADO INSTITUCIONAL — Franja superior tipo publicación
    # ═══════════════════════════════════════════════════════════

    # Barra institucional azul gobierno
    t_banner = doc.add_table(rows=1, cols=1)
    t_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_banner)
    banner_cell = t_banner.cell(0, 0)
    _shade_cell(banner_cell, _AZUL_GOB)
    banner_cell.text = ""
    p = banner_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  Servicio de Salud Metropolitano Central")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True
    run = p.add_run("  ·  ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0xBB, 0xDD)
    run = p.add_run("Ministerio de Salud — Gobierno de Chile")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xBB, 0xCC, 0xDD)

    # Línea roja institucional
    t_line = doc.add_table(rows=1, cols=1)
    t_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_line)
    line_cell = t_line.cell(0, 0)
    _shade_cell(line_cell, _ROJO_GOB)
    line_cell.text = ""
    p = line_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" ")
    run.font.size = Pt(2)

    _add_spacer(4)

    # ── TÍTULO PRINCIPAL — Estilo revista ──
    t_title = doc.add_table(rows=1, cols=1)
    t_title.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_title)
    title_cell = t_title.cell(0, 0)
    _shade_cell(title_cell, _VERDE_SSMC)
    title_cell.text = ""
    # Línea 1: ACTA DE REUNIÓN
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ACTA DE REUNIÓN")
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True
    run.font.all_caps = True
    # Línea 2: Subtítulo
    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("Dirección de Atención Primaria")
    run2.font.size = Pt(12)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x4A, 0xDE, 0x80)
    run2.bold = True
    # Línea 3: Detalle
    p3 = title_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(10)
    run3 = p3.add_run(f"Revisión de Dotación — {acta_cfg['cesfam']}")
    run3.font.size = Pt(10)
    run3.font.name = "Calibri"
    run3.font.color.rgb = RGBColor(0xBB, 0xCC, 0xBB)

    # Línea verde acento
    t_accent = doc.add_table(rows=1, cols=1)
    t_accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_accent)
    acc_cell = t_accent.cell(0, 0)
    _shade_cell(acc_cell, _VERDE_CLARO.replace("#", ""))
    acc_cell.text = ""
    p = acc_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" ")
    run.font.size = Pt(1)

    _add_spacer(8)

    # ═══════════════════════════════════════════════════════════
    # DATOS DE LA REUNIÓN — Tabla elegante
    # ═══════════════════════════════════════════════════════════
    t_head = doc.add_table(rows=3, cols=6)
    t_head.style = "Table Grid"
    t_head.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fila 0: Objetivo
    _set_cell(t_head.cell(0, 0), "Objetivo reunión", bold=True, size=9, color=_BLANCO)
    _shade_cell(t_head.cell(0, 0), _VERDE_MEDIO)
    t_head.cell(0, 1).merge(t_head.cell(0, 5))
    _set_cell(t_head.cell(0, 1), acta_cfg["objetivo"], size=9, bold=True, color=_VERDE_SSMC)
    _shade_cell(t_head.cell(0, 1), _VERDE_FONDO)

    # Fila 1: Fecha/hora inicio + lugar
    for c in range(6):
        _shade_cell(t_head.cell(1, c), _GRIS_CLARO)
    _set_cell(t_head.cell(1, 0), "Fecha inicio", bold=True, size=8, color=_GRIS_OSCURO)
    _set_cell(t_head.cell(1, 1), acta_cfg["fecha_inicio"], size=9)
    _set_cell(t_head.cell(1, 2), "Hora inicio", bold=True, size=8, color=_GRIS_OSCURO)
    _set_cell(t_head.cell(1, 3), acta_cfg["hora_inicio"], size=9)
    _set_cell(t_head.cell(1, 4), "Lugar", bold=True, size=8, color=_GRIS_OSCURO)
    _set_cell(t_head.cell(1, 5), acta_cfg["lugar"], size=9)

    # Fila 2: Fecha/hora fin
    _set_cell(t_head.cell(2, 0), "Fecha fin", bold=True, size=8, color=_GRIS_OSCURO)
    _set_cell(t_head.cell(2, 1), acta_cfg["fecha_fin"], size=9)
    _set_cell(t_head.cell(2, 2), "Hora fin", bold=True, size=8, color=_GRIS_OSCURO)
    _set_cell(t_head.cell(2, 3), acta_cfg["hora_fin"], size=9)
    t_head.cell(2, 4).merge(t_head.cell(2, 5))
    _set_cell(t_head.cell(2, 4), "", size=9)

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # NÓMINA DE PARTICIPANTES
    # ═══════════════════════════════════════════════════════════

    # Sección header bar
    t_sec1 = doc.add_table(rows=1, cols=1)
    t_sec1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec1)
    sec1_cell = t_sec1.cell(0, 0)
    _shade_cell(sec1_cell, _VERDE_SSMC)
    sec1_cell.text = ""
    p = sec1_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  NÓMINA DE PARTICIPANTES")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(2)

    participantes = acta_cfg.get("participantes", [])
    t_part = doc.add_table(rows=1 + max(len(participantes), 1), cols=2)
    t_part.style = "Table Grid"
    t_part.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    _set_cell(t_part.cell(0, 0), "Nombre", bold=True, size=9, color=_BLANCO)
    _shade_cell(t_part.cell(0, 0), _VERDE_MEDIO)
    _set_cell(t_part.cell(0, 1), "DPTO/UNIDAD/CESFAM", bold=True, size=9, color=_BLANCO)
    _shade_cell(t_part.cell(0, 1), _VERDE_MEDIO)

    for i, p_item in enumerate(participantes):
        bg = _GRIS_CLARO if i % 2 == 0 else _BLANCO
        _set_cell(t_part.cell(i + 1, 0), p_item.get("nombre", ""), size=9)
        _shade_cell(t_part.cell(i + 1, 0), bg)
        _set_cell(t_part.cell(i + 1, 1), p_item.get("cargo", ""), size=9)
        _shade_cell(t_part.cell(i + 1, 1), bg)

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # ANTECEDENTES DEL CENTRO
    # ═══════════════════════════════════════════════════════════
    cesfam = acta_cfg["cesfam"]

    t_sec2 = doc.add_table(rows=1, cols=1)
    t_sec2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec2)
    sec2_cell = t_sec2.cell(0, 0)
    _shade_cell(sec2_cell, _VERDE_SSMC)
    sec2_cell.text = ""
    p = sec2_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  ANTECEDENTES — {cesfam}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(4)

    # Info en caja elegante
    info_items = []
    if acta_data["sectores"]:
        info_items.append(("Sectores", acta_data['sectores']))
    if acta_data["poblacion"]:
        info_items.append(("Población inscrita", acta_data['poblacion']))
    if acta_data["cecosf"]:
        info_items.append(("CECOSF", acta_data['cecosf']))
    if acta_cfg.get("colores_sectores"):
        info_items.append(("Colores sectores", acta_cfg['colores_sectores']))
    sapu_val = "Sí" if acta_cfg.get("tiene_sapu") else "No"
    info_items.append(("Cuenta con SAPU", sapu_val))
    if acta_cfg.get("tiene_sapu") and acta_cfg.get("obs_sapu"):
        info_items.append(("Observaciones SAPU", acta_cfg['obs_sapu']))

    t_info = doc.add_table(rows=len(info_items), cols=2)
    t_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_info)
    for i, (label, value) in enumerate(info_items):
        bg = _VERDE_FONDO if i % 2 == 0 else _BLANCO
        _set_cell(t_info.cell(i, 0), f"  {label}:", bold=True, size=9, color=_VERDE_SSMC)
        _shade_cell(t_info.cell(i, 0), bg)
        _set_cell(t_info.cell(i, 1), value, size=9)
        _shade_cell(t_info.cell(i, 1), bg)

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # RESUMEN DOTACIÓN POR ESTAMENTO
    # ═══════════════════════════════════════════════════════════

    t_sec3 = doc.add_table(rows=1, cols=1)
    t_sec3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec3)
    sec3_cell = t_sec3.cell(0, 0)
    _shade_cell(sec3_cell, _VERDE_SSMC)
    sec3_cell.text = ""
    p = sec3_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  RESUMEN DOTACIÓN POR ESTAMENTO")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(2)

    est_data = acta_data["estamentos"]
    if est_data:
        cols_est = ["Estamento", "N° Funcionarios", "Hrs Totales", "Hrs Descuento",
                    "Hrs Clínicas", "Jornadas Reales", "Jornadas Ideales", "Brecha"]
        t_est = doc.add_table(rows=1 + len(est_data) + 1, cols=len(cols_est))
        t_est.style = "Table Grid"
        t_est.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Encabezados con verde oscuro
        for j, col_name in enumerate(cols_est):
            _set_cell(t_est.cell(0, j), col_name, bold=True, size=7, align="center", color=_BLANCO)
            _shade_cell(t_est.cell(0, j), _VERDE_MEDIO)

        # Datos con bandas alternadas
        sum_func = 0
        sum_ht = 0
        sum_hd = 0
        sum_hc = 0
        sum_jr = 0
        for i, row in enumerate(est_data):
            row_bg = _GRIS_CLARO if i % 2 == 0 else _BLANCO
            for j, col_name in enumerate(cols_est):
                val = row.get(col_name, "")
                align = "center" if j > 0 else "left"
                _set_cell(t_est.cell(i + 1, j), val, size=8, align=align)
                # Color brecha (sobreescribe banda)
                if col_name == "Brecha" and isinstance(val, (int, float)):
                    if val < -0.05:
                        _shade_cell(t_est.cell(i + 1, j), "FDE8E8")  # rojo suave
                        _set_cell(t_est.cell(i + 1, j), val, size=8, align=align, bold=True, color="C53030")
                    elif val > 0.05:
                        _shade_cell(t_est.cell(i + 1, j), "FEFCE8")  # amarillo suave
                        _set_cell(t_est.cell(i + 1, j), val, size=8, align=align, bold=True, color="975A16")
                    else:
                        _shade_cell(t_est.cell(i + 1, j), "E6F4EA")  # verde suave
                        _set_cell(t_est.cell(i + 1, j), val, size=8, align=align, bold=True, color="276749")
                elif col_name != "Brecha":
                    _shade_cell(t_est.cell(i + 1, j), row_bg)

            sum_func += row.get("N° Funcionarios", 0)
            sum_ht += row.get("Hrs Totales", 0)
            sum_hd += row.get("Hrs Descuento", 0)
            sum_hc += row.get("Hrs Clínicas", 0)
            sum_jr += row.get("Jornadas Reales", 0)

        # Fila TOTAL con estilo destacado
        last_row = len(est_data) + 1
        _set_cell(t_est.cell(last_row, 0), "TOTAL", bold=True, size=8, color=_BLANCO)
        _shade_cell(t_est.cell(last_row, 0), _VERDE_SSMC)
        totals = [sum_func, round(sum_ht, 1), round(sum_hd, 1), round(sum_hc, 1), round(sum_jr, 2), "", ""]
        for j, val in enumerate(totals):
            _set_cell(t_est.cell(last_row, j + 1), val, bold=True, size=8, align="center", color=_BLANCO)
            _shade_cell(t_est.cell(last_row, j + 1), _VERDE_SSMC)

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # HORAS CLÍNICAS POR UNIDAD DE DESEMPEÑO
    # ═══════════════════════════════════════════════════════════
    unidades = acta_data.get("unidades", [])
    if unidades:
        t_sec_ud = doc.add_table(rows=1, cols=1)
        t_sec_ud.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(t_sec_ud)
        sec_ud_cell = t_sec_ud.cell(0, 0)
        _shade_cell(sec_ud_cell, _VERDE_SSMC)
        sec_ud_cell.text = ""
        p = sec_ud_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("  HORAS CLÍNICAS POR UNIDAD DE DESEMPEÑO")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True

        _add_spacer(2)

        cols_ud = ["Unidad", "N° Func.", "Hrs Totales", "Hrs Descuento", "Hrs Clínicas", "Jornadas"]
        t_ud = doc.add_table(rows=1 + len(unidades) + 1, cols=len(cols_ud))
        t_ud.style = "Table Grid"
        t_ud.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Encabezados
        for j, col_name in enumerate(cols_ud):
            _set_cell(t_ud.cell(0, j), col_name, bold=True, size=7, align="center", color=_BLANCO)
            _shade_cell(t_ud.cell(0, j), _VERDE_MEDIO)

        # Datos con bandas alternadas
        sum_func_ud = 0
        sum_ht_ud = 0
        sum_hd_ud = 0
        sum_hc_ud = 0
        sum_jr_ud = 0
        for i, row_ud in enumerate(unidades):
            row_bg = _GRIS_CLARO if i % 2 == 0 else _BLANCO
            for j, col_name in enumerate(cols_ud):
                val = row_ud.get(col_name, "")
                align = "center" if j > 0 else "left"
                _set_cell(t_ud.cell(i + 1, j), val, size=8, align=align)
                _shade_cell(t_ud.cell(i + 1, j), row_bg)
            sum_func_ud += row_ud.get("N° Func.", 0)
            sum_ht_ud += row_ud.get("Hrs Totales", 0)
            sum_hd_ud += row_ud.get("Hrs Descuento", 0)
            sum_hc_ud += row_ud.get("Hrs Clínicas", 0)
            sum_jr_ud += row_ud.get("Jornadas", 0)

        # Fila TOTAL
        last_ud = len(unidades) + 1
        _set_cell(t_ud.cell(last_ud, 0), "TOTAL", bold=True, size=8, color=_BLANCO)
        _shade_cell(t_ud.cell(last_ud, 0), _VERDE_SSMC)
        totals_ud = [sum_func_ud, round(sum_ht_ud, 1), round(sum_hd_ud, 1),
                     round(sum_hc_ud, 1), round(sum_jr_ud, 2)]
        for j, val in enumerate(totals_ud):
            _set_cell(t_ud.cell(last_ud, j + 1), val, bold=True, size=8, align="center", color=_BLANCO)
            _shade_cell(t_ud.cell(last_ud, j + 1), _VERDE_SSMC)

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # DETALLE DE LA REUNIÓN — Observaciones y Acuerdos
    # ═══════════════════════════════════════════════════════════

    t_sec4 = doc.add_table(rows=1, cols=1)
    t_sec4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec4)
    sec4_cell = t_sec4.cell(0, 0)
    _shade_cell(sec4_cell, _VERDE_SSMC)
    sec4_cell.text = ""
    p = sec4_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  DETALLE DE LA REUNIÓN")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(4)

    obs_text = acta_cfg.get("observaciones", "")
    if obs_text:
        # Caja con borde verde para observaciones
        t_obs = doc.add_table(rows=1, cols=1)
        t_obs.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(t_obs)
        obs_cell = t_obs.cell(0, 0)
        _shade_cell(obs_cell, "F0FDF4")
        _set_cell_border(obs_cell,
                         left={"val": "single", "sz": "12", "color": _VERDE_CLARO.replace("#", "")},
                         top={"val": "single", "sz": "2", "color": "E2E8F0"},
                         bottom={"val": "single", "sz": "2", "color": "E2E8F0"},
                         right={"val": "single", "sz": "2", "color": "E2E8F0"})
        obs_cell.text = ""
        for line_idx, line in enumerate(obs_text.split("\n")):
            if line.strip():
                if line_idx == 0:
                    p = obs_cell.paragraphs[0]
                else:
                    p = obs_cell.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"  • {line.strip()}")
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    else:
        p = doc.add_paragraph()
        run = p.add_run("(Sin observaciones registradas)")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run.italic = True

    _add_spacer(10)

    # ═══════════════════════════════════════════════════════════
    # OBSERVACIONES GENERALES — Cuadro de cierre
    # ═══════════════════════════════════════════════════════════
    t_sec_og = doc.add_table(rows=1, cols=1)
    t_sec_og.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec_og)
    sec_og_cell = t_sec_og.cell(0, 0)
    _shade_cell(sec_og_cell, _VERDE_SSMC)
    sec_og_cell.text = ""
    p = sec_og_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("OBSERVACIONES GENERALES")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(2)

    # Caja de observaciones generales
    obs_gral_text = acta_cfg.get("obs_generales", "")
    t_og_box = doc.add_table(rows=1, cols=1)
    t_og_box.style = "Table Grid"
    t_og_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    og_cell = t_og_box.cell(0, 0)
    _shade_cell(og_cell, _VERDE_FONDO)
    og_cell.text = ""
    p = og_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if obs_gral_text.strip():
        run = p.add_run(obs_gral_text.strip())
    else:
        run = p.add_run("(Sin observaciones generales)")
        run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    _add_spacer(4)

    # Tabla: Elaborada por / Fecha / Próxima reunión / Lugar
    t_cierre = doc.add_table(rows=2, cols=4)
    t_cierre.style = "Table Grid"
    t_cierre.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_cell(t_cierre.cell(0, 0), "Acta elaborada por", bold=True, size=8, color=_VERDE_SSMC)
    _shade_cell(t_cierre.cell(0, 0), _GRIS_CLARO)
    _set_cell(t_cierre.cell(0, 1), acta_cfg.get("elaborada_por", ""), size=9)
    _set_cell(t_cierre.cell(0, 2), "Fecha", bold=True, size=8, color=_VERDE_SSMC)
    _shade_cell(t_cierre.cell(0, 2), _GRIS_CLARO)
    _set_cell(t_cierre.cell(0, 3), acta_cfg.get("fecha_elaboracion", ""), size=9)

    _set_cell(t_cierre.cell(1, 0), "Próxima Reunión:", bold=True, size=8, color=_VERDE_SSMC)
    _shade_cell(t_cierre.cell(1, 0), _GRIS_CLARO)
    _set_cell(t_cierre.cell(1, 1), acta_cfg.get("prox_reunion", "Sin fecha"), size=9)
    _set_cell(t_cierre.cell(1, 2), "Lugar:", bold=True, size=8, color=_VERDE_SSMC)
    _shade_cell(t_cierre.cell(1, 2), _GRIS_CLARO)
    _set_cell(t_cierre.cell(1, 3), acta_cfg.get("prox_lugar", "Por definir"), size=9)

    _add_spacer(4)

    # Nota importante (disclaimer)
    t_nota = doc.add_table(rows=1, cols=1)
    t_nota.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_nota)
    nota_cell = t_nota.cell(0, 0)
    _shade_cell(nota_cell, "FFF8E1")  # amarillo muy claro
    _set_cell_border(nota_cell,
                     left={"val": "single", "sz": "12", "color": "F59E0B"},
                     top={"val": "single", "sz": "2", "color": "FDE68A"},
                     bottom={"val": "single", "sz": "2", "color": "FDE68A"},
                     right={"val": "single", "sz": "2", "color": "FDE68A"})
    nota_cell.text = ""
    p = nota_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Importante:")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    run.bold = True
    run.italic = True
    p2 = nota_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(
        "Las observaciones respecto al contenido de la minuta, tanto para los temas tratados "
        "en la reunión correspondiente como para los compromisos registrados, deberán ser "
        "enviadas por mail a: constanza.calderonq@redsalud.gob.cl con copia a "
        "katherine.vera@redsalud.gob.cl en un plazo no menor a cinco días hábiles, "
        "de lo contrario se entenderá como aceptada."
    )
    run2.font.size = Pt(7)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    run2.italic = True

    _add_spacer(14)

    # ═══════════════════════════════════════════════════════════
    # FIRMAS — Diseño institucional
    # ═══════════════════════════════════════════════════════════

    t_sec5 = doc.add_table(rows=1, cols=1)
    t_sec5.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_sec5)
    sec5_cell = t_sec5.cell(0, 0)
    _shade_cell(sec5_cell, _VERDE_SSMC)
    sec5_cell.text = ""
    p = sec5_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  FIRMAS")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    _add_spacer(20)

    t_firma = doc.add_table(rows=3, cols=3)
    t_firma.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_firma)

    labels = ["Subdirección GDP", "Director(a) CESFAM", "Subdirección Técnica"]
    for j in range(3):
        # Línea de firma
        _set_cell(t_firma.cell(0, j), "", size=9, align="center")
        # Línea horizontal
        sep_cell = t_firma.cell(1, j)
        sep_cell.text = ""
        p = sep_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("_" * 28)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0A, 0x2E, 0x1F)
        # Título del cargo
        _set_cell(t_firma.cell(2, j), labels[j], bold=True, size=8, align="center", color=_VERDE_SSMC)

    _add_spacer(14)

    # ═══════════════════════════════════════════════════════════
    # PIE DE PÁGINA — Franja institucional
    # ═══════════════════════════════════════════════════════════
    t_footer_line = doc.add_table(rows=1, cols=1)
    t_footer_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_footer_line)
    fl_cell = t_footer_line.cell(0, 0)
    _shade_cell(fl_cell, _ROJO_GOB)
    fl_cell.text = ""
    p = fl_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" ")
    run.font.size = Pt(1)

    t_footer = doc.add_table(rows=1, cols=1)
    t_footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(t_footer)
    footer_cell = t_footer.cell(0, 0)
    _shade_cell(footer_cell, _AZUL_GOB)
    footer_cell.text = ""
    p = footer_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Servicio de Salud Metropolitano Central · Dirección de Atención Primaria")
    run.font.size = Pt(7)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xBB, 0xCC, 0xDD)
    run = p.add_run("\n")
    run.font.size = Pt(4)
    run = p.add_run("Merced #280, Santiago · www.ssmc.cl")
    run.font.size = Pt(7)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x88, 0xAA, 0xCC)

    # Guardar a bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def page_acta(filtros):
    st.markdown("""
    <div class="main-header">
        <h1>📝 Acta de Reunión — Revisión de Dotación</h1>
        <p>Generación automática del acta con datos calculados por estamento y CESFAM</p>
    </div>
    """, unsafe_allow_html=True)

    df = st.session_state.df_main
    df_dot = st.session_state.get("df_dot", None)

    if df is None:
        st.warning("Cargue datos primero.")
        return

    # ── Paso 1: Selección de CESFAM ──
    st.markdown('<div class="section-title">1. Seleccionar Centro de Salud</div>', unsafe_allow_html=True)
    cesfam_list = sorted(df["CESFAM"].dropna().unique().tolist()) if "CESFAM" in df.columns else []
    if not cesfam_list:
        st.error("No hay CESFAM disponibles en los datos.")
        return

    sel_cesfam = st.selectbox("CESFAM", cesfam_list, key="acta_cesfam",
                               placeholder="Seleccione un centro de salud...")
    if not sel_cesfam:
        st.info("Seleccione un CESFAM para generar el acta.")
        return

    # Calcular datos automáticos
    acta_data = _build_acta_data(sel_cesfam, df, df_dot)

    # Mostrar resumen calculado
    st.markdown(f"""<div style="background:linear-gradient(135deg,#0A2E1F,#145A38);
        color:#fff; padding:.8rem 1.2rem; border-radius:8px; margin:.5rem 0;
        border-left:5px solid #4ADE80;">
        <strong>📊 {sel_cesfam}</strong> — {acta_data['total_personas']} funcionarios ·
        Sectores: {acta_data['sectores'] or '—'} · Población: {acta_data['poblacion'] or '—'} · CECOSF: {acta_data['cecosf'] or '—'}
    </div>""", unsafe_allow_html=True)

    # Tabla resumen estamentos
    if acta_data["estamentos"]:
        df_est = pd.DataFrame(acta_data["estamentos"])
        # Convertir "—" a NaN para display limpio
        df_est_display = df_est.copy()
        for c in ["Jornadas Ideales", "Brecha"]:
            if c in df_est_display.columns:
                df_est_display[c] = pd.to_numeric(df_est_display[c], errors="coerce")
        st.dataframe(df_est_display, use_container_width=True, hide_index=True, height=min(400, 40 + 35 * len(df_est)))
    else:
        st.info("No se encontraron estamentos con datos para este CESFAM.")

    # Tabla resumen por unidad de desempeño
    if acta_data.get("unidades"):
        st.markdown('<div class="section-title">Horas Clínicas por Unidad de Desempeño</div>', unsafe_allow_html=True)
        df_ud = pd.DataFrame(acta_data["unidades"])
        st.dataframe(df_ud, use_container_width=True, hide_index=True, height=min(400, 40 + 35 * len(df_ud)))

    st.divider()

    # ── Paso 2: Datos editables del acta ──
    st.markdown('<div class="section-title">2. Datos de la Reunión (editables)</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        objetivo = st.text_input("Objetivo de la reunión",
                                  value=f"Revisión Dotación {sel_cesfam}",
                                  key="acta_objetivo")
        fecha_inicio = st.date_input("Fecha inicio", value=date.today(), key="acta_fecha_ini")
        hora_inicio = st.time_input("Hora inicio", value=dt_time(13, 0), key="acta_hora_ini")
    with c2:
        lugar = st.text_input("Lugar", value="Merced #280, piso 10", key="acta_lugar")
        fecha_fin = st.date_input("Fecha fin", value=date.today(), key="acta_fecha_fin")
        hora_fin = st.time_input("Hora fin", value=dt_time(17, 0), key="acta_hora_fin")

    st.divider()

    # ── Paso 3: Información variable del centro (sectores, SAPU) ──
    st.markdown('<div class="section-title">3. Información del Centro (variable por CESFAM)</div>', unsafe_allow_html=True)

    c3, c4 = st.columns(2)
    with c3:
        colores_sectores = st.text_input(
            "Colores de sectores (ej: Rojo, Azul, Verde, Naranjo)",
            value="", key="acta_colores",
            placeholder="Ingrese los colores separados por coma...")
        tiene_sapu = st.checkbox("¿Tiene SAPU?", value=False, key="acta_sapu")
    with c4:
        if tiene_sapu:
            obs_sapu = st.text_input("Observaciones SAPU",
                                      value="", key="acta_obs_sapu",
                                      placeholder="Horario, dotación SAPU, etc.")
        else:
            obs_sapu = ""

    st.divider()

    # ── Paso 4: Nómina de participantes ──
    st.markdown('<div class="section-title">4. Nómina de Participantes</div>', unsafe_allow_html=True)

    # Inicializar participantes en session_state (lista de dicts)
    if "acta_part_list" not in st.session_state:
        st.session_state.acta_part_list = [{"nombre": "", "unidad": ""} for _ in range(5)]

    n_part = len(st.session_state.acta_part_list)

    # Renderizar filas con text_input individuales (sin lag de data_editor)
    for idx in range(n_part):
        c_n, c_u = st.columns(2)
        with c_n:
            val_n = st.text_input(
                f"Nombre #{idx+1}", value=st.session_state.acta_part_list[idx]["nombre"],
                key=f"_ap_n_{idx}", placeholder="Nombre del participante",
                label_visibility="collapsed" if idx > 0 else "visible",
            )
            st.session_state.acta_part_list[idx]["nombre"] = val_n
        with c_u:
            val_u = st.text_input(
                f"DPTO/UNIDAD/CESFAM #{idx+1}", value=st.session_state.acta_part_list[idx]["unidad"],
                key=f"_ap_u_{idx}", placeholder="Depto / Unidad / CESFAM",
                label_visibility="collapsed" if idx > 0 else "visible",
            )
            st.session_state.acta_part_list[idx]["unidad"] = val_u

    bc1, bc2, _ = st.columns([1, 1, 4])
    with bc1:
        if st.button("➕ Agregar fila", key="acta_add_part"):
            st.session_state.acta_part_list.append({"nombre": "", "unidad": ""})
            st.rerun()
    with bc2:
        if n_part > 1 and st.button("➖ Quitar última", key="acta_rm_part"):
            st.session_state.acta_part_list.pop()
            st.rerun()

    st.divider()

    # ── Paso 5: Observaciones y acuerdos ──
    st.markdown('<div class="section-title">5. Observaciones y Acuerdos</div>', unsafe_allow_html=True)

    observaciones = st.text_area(
        "Observaciones y acuerdos de la reunión",
        value="", key="acta_obs",
        height=150,
        placeholder="Ingrese observaciones, acuerdos, compromisos...",
    )

    st.divider()

    # ── Paso 6: Observaciones Generales y Cierre ──
    st.markdown('<div class="section-title">6. Observaciones Generales y Cierre</div>', unsafe_allow_html=True)

    obs_generales = st.text_area(
        "Observaciones generales",
        value="", key="acta_obs_gral",
        height=100,
        placeholder="Observaciones generales de la reunión...",
    )

    c5, c6 = st.columns(2)
    with c5:
        elaborada_por = st.text_input("Acta elaborada por",
                                       value="", key="acta_elaborada",
                                       placeholder="Nombre completo")
        prox_reunion = st.text_input("Próxima reunión (fecha)",
                                      value="Sin fecha", key="acta_prox_fecha")
    with c6:
        fecha_elaboracion = st.text_input("Fecha elaboración",
                                           value=date.today().strftime("%m-%Y"),
                                           key="acta_fecha_elab")
        prox_lugar = st.text_input("Lugar próxima reunión",
                                    value="Por definir", key="acta_prox_lugar")

    st.divider()

    # ── Paso 7: Generar y descargar ──
    st.markdown('<div class="section-title">7. Generar Acta</div>', unsafe_allow_html=True)

    # Preparar participantes (filtrar vacíos)
    participantes = []
    for p_item in st.session_state.get("acta_part_list", []):
        nombre = str(p_item.get("nombre", "")).strip()
        cargo = str(p_item.get("unidad", "")).strip()
        if nombre:
            participantes.append({"nombre": nombre, "cargo": cargo})

    acta_cfg = {
        "cesfam": sel_cesfam,
        "objetivo": objetivo,
        "fecha_inicio": fecha_inicio.strftime("%d-%m-%Y"),
        "fecha_fin": fecha_fin.strftime("%d-%m-%Y"),
        "hora_inicio": hora_inicio.strftime("%H:%M"),
        "hora_fin": hora_fin.strftime("%H:%M"),
        "lugar": lugar,
        "participantes": participantes,
        "colores_sectores": colores_sectores,
        "tiene_sapu": tiene_sapu,
        "obs_sapu": obs_sapu,
        "observaciones": observaciones,
        "obs_generales": obs_generales,
        "elaborada_por": elaborada_por,
        "fecha_elaboracion": fecha_elaboracion,
        "prox_reunion": prox_reunion,
        "prox_lugar": prox_lugar,
    }

    # Vista previa
    with st.expander("👀 Vista previa del acta", expanded=False):
        st.markdown(f"**Objetivo:** {objetivo}")
        st.markdown(f"**Fecha:** {acta_cfg['fecha_inicio']} · {acta_cfg['hora_inicio']} – {acta_cfg['hora_fin']}")
        st.markdown(f"**Lugar:** {lugar}")
        if participantes:
            st.markdown("**Participantes:**")
            for p in participantes:
                st.markdown(f"- {p['nombre']} — {p['cargo']}")
        if colores_sectores:
            st.markdown(f"**Sectores:** {colores_sectores}")
        st.markdown(f"**SAPU:** {'Sí' if tiene_sapu else 'No'}")
        if acta_data["estamentos"]:
            st.markdown("**Resumen por estamento:**")
            _prev_df = pd.DataFrame(acta_data["estamentos"])
            for _c in ["Jornadas Ideales", "Brecha"]:
                if _c in _prev_df.columns:
                    _prev_df[_c] = pd.to_numeric(_prev_df[_c], errors="coerce")
            st.dataframe(_prev_df, use_container_width=True, hide_index=True)

    # Botón generar
    col_btn, col_info = st.columns([1, 3])
    with col_btn:
        generate = st.button("📄 Generar Acta Word", type="primary", use_container_width=True,
                              key="acta_generate")
    with col_info:
        st.caption("Se generará un documento Word (.docx) con todos los datos calculados y editables.")

    if generate:
        with st.spinner("Generando acta..."):
            docx_bytes = _generate_acta_docx(acta_cfg, acta_data)
            safe_name = sel_cesfam.replace(" ", "_").replace("/", "-")
            fname = f"Acta_Reunion_{safe_name}_{fecha_inicio.strftime('%Y%m%d')}.docx"
            st.session_state["_acta_bytes"] = docx_bytes
            st.session_state["_acta_fname"] = fname
            st.toast("Acta generada correctamente", icon="✅")

    # Botón de descarga (aparece después de generar)
    if "_acta_bytes" in st.session_state:
        st.download_button(
            "⬇️ Descargar Acta",
            data=st.session_state["_acta_bytes"],
            file_name=st.session_state["_acta_fname"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )


# ─────────────────────────────────────────────────────────────
# AUTENTICACIÓN SIMPLE
# ─────────────────────────────────────────────────────────────
def _check_password() -> bool:
    """Muestra login y retorna True si la contraseña es correcta."""
    _pwd = st.secrets.get("password", "dap2026")  # fallback para desarrollo local

    def _verify():
        if hmac.compare_digest(st.session_state.get("_pwd_input", ""), _pwd):
            st.session_state["_authenticated"] = True
            del st.session_state["_pwd_input"]
        else:
            st.session_state["_authenticated"] = False

    if st.session_state.get("_authenticated"):
        return True

    st.markdown("""
    <div style="display:flex;justify-content:center;align-items:center;min-height:60vh;">
        <div style="text-align:center;">
            <h1 style="color:#0A2E1F;">🏥 Sistema de Gestión de Dotación APS</h1>
            <p style="color:#64748B;">Servicio de Salud Metropolitano Central</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.text_input("Contraseña", type="password", key="_pwd_input", on_change=_verify)
        if "_authenticated" in st.session_state and not st.session_state["_authenticated"]:
            st.error("Contraseña incorrecta")
    return False


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────
def main():
    if not _check_password():
        return

    nav, filtros = render_sidebar()

    if nav == "📊 Panel Principal":
        page_panel(filtros)
    elif nav == "📋 Dotación (Editor)":
        page_editor(filtros)
    elif nav == "⏱️ Horas Indirectas":
        page_horas()
    elif nav == "🎯 Dotación Ideal":
        page_dot_ideal()
    elif nav == "📈 Análisis de Brechas":
        page_brechas(filtros)
    elif nav == "📝 Acta de Reunión":
        page_acta(filtros)


if __name__ == "__main__":
    main()
